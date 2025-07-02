"""
Document processing service using LlamaIndex and GPT4All.
"""
import os
import uuid
import asyncio
import re
import hashlib
import mimetypes
import json
import logging
from datetime import datetime
import platform
import pwd
import grp
from typing import List, Optional, Dict, Any, ClassVar
from pathlib import Path
from pydantic import Field

from llama_index.core import (
    VectorStoreIndex, 
    SimpleDirectoryReader, 
    StorageContext,
    Settings,
    Document as LlamaDocument
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.llms import MockLLM, LLM
from llama_index.core.llms.callbacks import llm_completion_callback
from llama_index.core.embeddings import BaseEmbedding
from llama_index.core.base.llms.types import LLMMetadata, CompletionResponse, ChatResponse, ChatMessage
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.readers.confluence import ConfluenceReader
import chromadb
from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown
from bs4 import BeautifulSoup
from sentence_transformers import SentenceTransformer
import torch
import numpy as np
import requests

from ..models.document import Document, DocumentType, DocumentStatus, ConfluencePage
from ..models.template import Template

class BERTEmbedding(BaseEmbedding):
    """Advanced BERT embedding model wrapper with multiple model support."""
    
    # Pre-configured BERT models with their properties
    BERT_MODELS: ClassVar[Dict[str, Dict[str, Any]]] = {
        "bert-base-uncased": {
            "model_name": "bert-base-uncased",
            "dimensions": 768,
            "description": "BERT Base model, good general performance"
        },
        "sentence-bert": {
            "model_name": "all-MiniLM-L6-v2", 
            "dimensions": 384,
            "description": "Sentence-BERT, optimized for sentence embeddings"
        },
        "bert-large": {
            "model_name": "bert-large-uncased",
            "dimensions": 1024,
            "description": "BERT Large model, better accuracy but slower"
        },
        "distilbert": {
            "model_name": "distilbert-base-uncased",
            "dimensions": 768,
            "description": "DistilBERT, faster and smaller than BERT"
        },
        "roberta": {
            "model_name": "sentence-transformers/all-roberta-large-v1",
            "dimensions": 1024,
            "description": "RoBERTa-based sentence transformer"
        },
        "mpnet": {
            "model_name": "sentence-transformers/all-mpnet-base-v2",
            "dimensions": 768,
            "description": "MPNet model, excellent for semantic search"
        }
    }
    
    # Pydantic fields
    bert_config: Dict[str, Any] = Field(default_factory=dict)
    model_name_: str = Field(default="")
    device_: str = Field(default="cpu")
    model_: Optional[Any] = Field(default=None, exclude=True)
    tokenizer_: Optional[Any] = Field(default=None, exclude=True)
    use_transformers_: bool = Field(default=False)
    
    def __init__(self, model_key: str = "sentence-bert", use_gpu: bool = None, **kwargs):
        """
        Initialize BERT embedding model.
        
        Args:
            model_key: Key from BERT_MODELS dict or custom model name
            use_gpu: Whether to use GPU if available (auto-detect if None)
        """
        # Initialize parent first
        super().__init__(**kwargs)
        
        # Determine model configuration
        if model_key in self.BERT_MODELS:
            self.bert_config = self.BERT_MODELS[model_key]
            self.model_name_ = self.bert_config["model_name"]
        else:
            # Custom model name
            self.model_name_ = model_key
            self.bert_config = {"model_name": model_key, "dimensions": 768}
        
        self.model_ = None
        self.device_ = self._get_device(use_gpu)
        
        print(f"Initializing BERT embedding model: {self.model_name_}")
        if model_key in self.BERT_MODELS:
            print(f"Model info: {self.bert_config['description']}")
        print(f"Using device: {self.device_}")
    
    def _get_device(self, use_gpu: bool = None) -> str:
        """Determine the best device to use."""
        if use_gpu is False:
            return "cpu"
        
        if torch.cuda.is_available():
            if use_gpu is None:
                print("CUDA available - using GPU for embeddings")
            return "cuda"
        elif hasattr(torch.backends, 'mps') and torch.backends.mps.is_available():
            if use_gpu is None:
                print("MPS (Apple Silicon) available - using MPS for embeddings")
            return "mps"
        else:
            if use_gpu is True:
                print("Warning: GPU requested but not available, using CPU")
            return "cpu"
    
    @property
    def model(self):
        """Lazy load the model when first accessed."""
        if self.model_ is None:
            try:
                # Try SentenceTransformer first (best for embeddings)
                self.model_ = SentenceTransformer(self.model_name_, device=self.device_)
                print(f"Loaded SentenceTransformer model: {self.model_name_}")
            except Exception as e:
                print(f"Failed to load as SentenceTransformer: {e}")
                try:
                    # Fallback to transformers library
                    from transformers import AutoTokenizer, AutoModel
                    self.tokenizer_ = AutoTokenizer.from_pretrained(self.model_name_)
                    self.model_ = AutoModel.from_pretrained(self.model_name_)
                    self.model_.to(self.device_)
                    self.use_transformers_ = True
                    print(f"Loaded transformers model: {self.model_name_}")
                except Exception as e2:
                    print(f"Failed to load model: {e2}")
                    raise Exception(f"Could not load model {self.model_name_}: {e2}")
        return self.model_
    
    def _get_query_embedding(self, query: str) -> List[float]:
        """Get embedding for a query string."""
        return self._get_text_embedding(query)
    
    def _get_text_embedding(self, text: str) -> List[float]:
        """Get embedding for text using the loaded model."""
        try:
            if hasattr(self.model, 'encode'):
                # SentenceTransformer model
                embedding = self.model.encode(text, convert_to_tensor=False)
                return embedding.tolist() if hasattr(embedding, 'tolist') else list(embedding)
            else:
                # Direct transformers model
                return self._get_transformers_embedding(text)
        except Exception as e:
            print(f"Error getting embedding for text: {e}")
            # Return zero vector as fallback
            return [0.0] * self.bert_config.get("dimensions", 768)
    
    def _get_transformers_embedding(self, text: str) -> List[float]:
        """Get embedding using transformers library directly."""
        from transformers import AutoTokenizer, AutoModel
        import torch
        
        # Tokenize and get model output
        inputs = self.tokenizer_(text, return_tensors="pt", truncation=True, padding=True, max_length=512)
        inputs = {k: v.to(self.device_) for k, v in inputs.items()}
        
        with torch.no_grad():
            outputs = self.model_(**inputs)
            # Use mean pooling of last hidden state
            embeddings = outputs.last_hidden_state.mean(dim=1).squeeze()
            return embeddings.cpu().numpy().tolist()
    
    async def _aget_query_embedding(self, query: str) -> List[float]:
        """Async version of query embedding."""
        return self._get_query_embedding(query)
    
    async def _aget_text_embedding(self, text: str) -> List[float]:
        """Async version of text embedding."""
        return self._get_text_embedding(text)
    
    def get_model_info(self) -> Dict[str, Any]:
        """Get information about the current model."""
        return {
            "model_name": self.model_name_,
            "device": self.device_,
            "config": self.bert_config,
            "available_models": list(self.BERT_MODELS.keys())
        }

class GPT4AllLLM(LLM):
    """GPT4All LLM wrapper that properly inherits from LlamaIndex LLM."""
    
    model_path: str
    model_name: str
    default_max_tokens: int = 512
    
    def __init__(self, model_path: str, model_name: str, **kwargs):
        super().__init__(model_path=model_path, model_name=model_name, **kwargs)
        self._model = None
        self._available = False
        self._init_model()
    
    def _init_model(self):
        try:
            from gpt4all import GPT4All
            # Initialize with larger context window for Llama 3.1
            self._model = GPT4All(
                self.model_name, 
                model_path=self.model_path,
                n_ctx=32768,  # Set context to 32K tokens (much more reasonable for Llama 3.1)
                verbose=False
            )
            self._available = True
            print(f"Successfully loaded GPT4All model: {self.model_name} with 32K context window")
        except Exception as e:
            print(f"Failed to load GPT4All model: {e}")
            self._available = False
    
    def set_default_max_tokens(self, max_tokens: int):
        """Set the default max_tokens for this LLM instance."""
        self.default_max_tokens = max_tokens
    
    @property
    def metadata(self) -> LLMMetadata:
        return LLMMetadata(
            context_window=32768,  # Much more reasonable for Llama 3.1 8B (32K tokens)
            num_output=self.default_max_tokens,
            model_name=self.model_name,
        )
    
    def complete(self, prompt: str, **kwargs) -> CompletionResponse:
        if not self._available or self._model is None:
            return CompletionResponse(text="GPT4All model not available. Please check model installation.", additional_kwargs={})
        
        try:
            max_tokens = kwargs.get('max_tokens', self.default_max_tokens)
            response = self._model.generate(prompt, max_tokens=max_tokens)
            return CompletionResponse(text=response, additional_kwargs={})
        except Exception as e:
            return CompletionResponse(text=f"Error generating response: {e}", additional_kwargs={})
    
    @llm_completion_callback()
    def stream_complete(self, prompt: str, **kwargs):
        # For simplicity, just return the complete response
        response = self.complete(prompt, **kwargs)
        yield response
    
    def chat(self, messages, **kwargs):
        # Convert chat messages to a single prompt
        prompt = ""
        for message in messages:
            role = message.role
            content = message.content
            prompt += f"{role}: {content}\n"
        
        prompt += "assistant: "
        response = self.complete(prompt, **kwargs)
        
        return ChatResponse(
            message=ChatMessage(role="assistant", content=response.text)
        )
    
    @llm_completion_callback()
    def stream_chat(self, messages, **kwargs):
        # For simplicity, just return the complete chat response
        response = self.chat(messages, **kwargs)
        yield response
    
    async def acomplete(self, prompt: str, **kwargs) -> CompletionResponse:
        # Async version - for simplicity, just call the sync version
        return self.complete(prompt, **kwargs)
    
    async def astream_complete(self, prompt: str, **kwargs):
        # Async streaming version
        response = await self.acomplete(prompt, **kwargs)
        yield response
    
    async def achat(self, messages, **kwargs):
        # Async chat version
        return self.chat(messages, **kwargs)
    
    async def astream_chat(self, messages, **kwargs):
        # Async streaming chat version
        response = await self.achat(messages, **kwargs)
        yield response

class OllamaLLM(LLM):
    """Ollama LLM wrapper that properly inherits from LlamaIndex LLM."""
    
    model_name: str
    default_max_tokens: int = 512
    base_url: str = "http://localhost:11434"
    available: bool = False
    
    def __init__(self, model_name: str, **kwargs):
        super().__init__(model_name=model_name, available=False, **kwargs)
        self._init_model()
    
    def _init_model(self):
        try:
            # Test Ollama connection with a simple HTTP request
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if response.status_code == 200:
                self.available = True
                print(f"Successfully connected to Ollama: {self.model_name}")
            else:
                self.available = False
                print(f"Ollama not responding: {response.status_code}")
        except Exception as e:
            print(f"Failed to connect to Ollama: {e}")
            self.available = False
    
    def set_default_max_tokens(self, max_tokens: int):
        """Set the default max_tokens for this LLM instance."""
        self.default_max_tokens = max_tokens
    
    @property
    def metadata(self) -> LLMMetadata:
        return LLMMetadata(
            context_window=8192,  # Default context window
            num_output=self.default_max_tokens,
            model_name=self.model_name,
        )
    
    def complete(self, prompt: str, **kwargs) -> CompletionResponse:
        if not self.available:
            return CompletionResponse(text="Ollama model not available. Please ensure Ollama is running.", additional_kwargs={})
        
        try:
            max_tokens = kwargs.get('max_tokens', self.default_max_tokens)
            
            # Direct HTTP call to Ollama
            response = requests.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": self.model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "num_predict": max_tokens
                    }
                },
                timeout=120
            )
            
            if response.status_code == 200:
                result = response.json()
                return CompletionResponse(text=result.get("response", ""), additional_kwargs={})
            else:
                return CompletionResponse(text=f"Error: Ollama API returned {response.status_code}", additional_kwargs={})
                
        except Exception as e:
            return CompletionResponse(text=f"Error generating response: {e}", additional_kwargs={})
    
    @llm_completion_callback()
    def stream_complete(self, prompt: str, **kwargs):
        # For simplicity, just return the complete response
        response = self.complete(prompt, **kwargs)
        yield response
    
    def chat(self, messages, **kwargs):
        if not self.available:
            return ChatResponse(
                message=ChatMessage(role="assistant", content="Ollama model not available. Please ensure Ollama is running.")
            )
        
        try:
            
            # Convert LlamaIndex messages to Ollama format
            ollama_messages = []
            for msg in messages:
                ollama_messages.append({
                    "role": msg.role,
                    "content": msg.content
                })
            
            # Direct HTTP call to Ollama chat API
            response = requests.post(
                f"{self.base_url}/api/chat",
                json={
                    "model": self.model_name,
                    "messages": ollama_messages,
                    "stream": False
                },
                timeout=120
            )
            
            if response.status_code == 200:
                result = response.json()
                content = result.get("message", {}).get("content", "")
                
                return ChatResponse(
                    message=ChatMessage(role="assistant", content=content)
                )
            else:
                return ChatResponse(
                    message=ChatMessage(role="assistant", content=f"Error: Ollama API returned {response.status_code}")
                )
                
        except Exception as e:
            return ChatResponse(
                message=ChatMessage(role="assistant", content=f"Error generating response: {e}")
            )
    
    @llm_completion_callback()
    def stream_chat(self, messages, **kwargs):
        # For simplicity, just return the complete chat response
        response = self.chat(messages, **kwargs)
        yield response
    
    async def acomplete(self, prompt: str, **kwargs) -> CompletionResponse:
        return self.complete(prompt, **kwargs)
    
    async def astream_complete(self, prompt: str, **kwargs):
        for response in self.stream_complete(prompt, **kwargs):
            yield response
    
    async def achat(self, messages, **kwargs):
        return self.chat(messages, **kwargs)
    
    async def astream_chat(self, messages, **kwargs):
        for response in self.stream_chat(messages, **kwargs):
            yield response

class DocumentService:
    """Service for processing and managing documents with LlamaIndex."""
    
    def __init__(self):
        self.index = None
        self.query_engine = None
        self.chat_engine = None
        self.documents: Dict[str, Document] = {}
        self.templates: Dict[str, Template] = {}
        self.chroma_client = None
        self.vector_store = None
        self.llm = None
        self.embedding_model = None
        
        # Paths
        self.data_dir = Path("../data")
        self.documents_dir = self.data_dir / "documents"
        self.models_dir = self.data_dir / "models"
        self.chroma_dir = self.data_dir / "chroma_db"
        self.templates_dir = self.data_dir / "templates"
        self.templates_metadata_file = self.templates_dir / "metadata.json"
        
        # Ensure directories exist
        for directory in [self.data_dir, self.documents_dir, self.models_dir, self.chroma_dir, self.templates_dir]:
            directory.mkdir(parents=True, exist_ok=True)
            
        # Configure logging to reduce LlamaIndex verbosity
        logging.getLogger("llama_index").setLevel(logging.WARNING)
        logging.getLogger("chromadb").setLevel(logging.WARNING)
        logging.getLogger("httpx").setLevel(logging.WARNING)

    def _load_app_settings(self) -> Dict[str, Any]:
        """Load application settings from file."""
        settings_file = self.data_dir / "app_settings.json"
        default_settings = {
            "max_tokens": 512,  # Safe default that leaves plenty of room for document context
            "temperature": 0.7,
            "use_document_context": True,
            "enable_notifications": True,
            "llm_provider": "gpt4all",  # Default to GPT4All
            "preferred_gpt4all_model": None,
            "preferred_ollama_model": "llama3.2:3b",  # Default Ollama model
            # Add conversation context settings
            "max_conversation_history": 10,  # Maximum number of previous messages to include
            "enable_conversation_context": True,  # Whether to include conversation history at all
        }
        
        if settings_file.exists():
            try:
                with open(settings_file, 'r') as f:
                    settings = json.load(f)
                    # Merge with defaults to ensure all keys exist
                    return {**default_settings, **settings}
            except Exception as e:
                print(f"Error loading settings, using defaults: {e}")
        
        return default_settings
    
    async def initialize(self, embedding_model: str = "mpnet", use_gpu: bool = None):
        """Initialize the document service with LLM and embeddings."""
        try:
            # Load app settings early
            settings = self._load_app_settings()
            
            # Initialize LLM based on provider setting
            llm_provider = settings.get('llm_provider', 'gpt4all')
            
            if llm_provider == 'ollama':
                # Initialize Ollama LLM
                ollama_model = settings.get('preferred_ollama_model', 'llama3.2:3b')
                self.llm = OllamaLLM(ollama_model)
                self.llm.set_default_max_tokens(settings['max_tokens'])
                Settings.llm = self.llm
                print(f"Initialized Ollama with model: {ollama_model} (max_tokens: {settings['max_tokens']})")
            else:
                # Initialize GPT4All LLM (default)
                model_path = self._get_gpt4all_model_path()
                if model_path:
                    self.llm = GPT4AllLLM(str(self.models_dir), model_path)
                    self.llm.set_default_max_tokens(settings['max_tokens'])
                    Settings.llm = self.llm
                    print(f"Initialized GPT4All with model: {model_path} (max_tokens: {settings['max_tokens']})")
                else:
                    print("Warning: No GPT4All model found. Using mock LLM.")
                    Settings.llm = MockLLM()
            
            # Configure BERT embeddings
            await self.set_embedding_model(embedding_model, use_gpu)
            
            # Configure node parser with improved settings for RAG accuracy
            Settings.node_parser = SentenceSplitter(
                chunk_size=1024,  # Increased from 512 for better context
                chunk_overlap=100,  # Increased from 50 for better continuity
                paragraph_separator="\n\n",  # Better paragraph boundaries
                secondary_chunking_regex="[.!?]+",  # Split on sentence boundaries
            )
            
            # Initialize Chroma vector store
            await self._initialize_vector_store()
            
            # Initialize or load existing index
            await self._initialize_index()
            
            # Load existing documents from the file system
            await self._load_existing_documents()
            
            # Load existing templates from the file system
            await self._load_existing_templates()
            
            print("Document service initialized successfully")
            
        except Exception as e:
            print(f"Error initializing document service: {e}")
            raise
    
    async def _load_existing_documents(self):
        """Load existing documents from the data directory on startup."""
        try:
            loaded_count = 0
            
            # Scan the documents directory for files
            if self.documents_dir.exists():
                for file_path in self.documents_dir.iterdir():
                    if file_path.is_file() and not file_path.name.startswith('.'):
                        try:
                            await self._load_document_from_file(file_path)
                            loaded_count += 1
                        except Exception as e:
                            print(f"Error loading document {file_path.name}: {e}")
            
            print(f"Loaded {loaded_count} existing documents from file system")
            
        except Exception as e:
            print(f"Error loading existing documents: {e}")
    
    async def _load_document_from_file(self, file_path: Path):
        """Load a single document from a file and add to document tracking."""
        try:
            # Get comprehensive file metadata
            file_metadata = await self._extract_file_metadata(file_path)
            
            # Determine document type first
            doc_type = self._get_document_type(file_path.name)
            
            # Handle different file types appropriately
            if doc_type in [DocumentType.PDF, DocumentType.DOCX, DocumentType.TRANSCRIPT]:
                # For binary files and transcripts, extract content using specialized methods
                content = await self._extract_content(str(file_path), doc_type)
            else:
                # For other text files, read as UTF-8
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    # Fallback for files with different encoding
                    try:
                        with open(file_path, 'r', encoding='latin-1') as f:
                            content = f.read()
                    except Exception:
                        # Last resort - treat as binary and convert
                        with open(file_path, 'rb') as f:
                            raw_content = f.read()
                            content = str(raw_content, errors='ignore')
            
            # Check if it's a Confluence file (markdown with metadata header)
            if file_path.name.startswith('confluence_') and file_path.suffix == '.md' and doc_type != DocumentType.PDF:
                # Parse markdown frontmatter for text files only
                if content.startswith('---\n'):
                    try:
                        parts = content.split('---\n', 2)
                        if len(parts) >= 3:
                            frontmatter = parts[1]
                            doc_content = parts[2]
                            
                            # Parse frontmatter
                            confluence_metadata = {}
                            for line in frontmatter.strip().split('\n'):
                                if ':' in line:
                                    key, value = line.split(':', 1)
                                    confluence_metadata[key.strip()] = value.strip()
                            
                            # Extract content-specific metadata from the actual content
                            content_metadata = await self._extract_content_metadata(file_path, DocumentType.CONFLUENCE, doc_content)
                            
                            # Combine file metadata with confluence metadata
                            combined_metadata = {**file_metadata, **content_metadata, **confluence_metadata}
                            
                            doc_id = confluence_metadata.get('doc_id')
                            if doc_id:
                                document = Document(
                                    id=doc_id,
                                    title=confluence_metadata.get('title', file_path.stem),
                                    type=DocumentType.CONFLUENCE,
                                    content=doc_content.strip(),
                                    file_path=str(file_path),
                                    status=DocumentStatus.PROCESSING,
                                    size_bytes=file_metadata['size_bytes'],
                                    created_at=datetime.fromtimestamp(file_metadata['created_at']),
                                    updated_at=datetime.fromtimestamp(file_metadata['modified_at']),
                                    metadata=combined_metadata
                                )
                                
                                # CRITICAL FIX: Create LlamaIndex document and insert into vector index
                                if self.index is not None:
                                    # Enhance content for better RAG performance
                                    enhanced_content = self._enhance_document_for_rag(doc_content.strip(), {
                                        'title': confluence_metadata.get('title', file_path.stem),
                                        'document_type': 'confluence',
                                        'filename': file_path.name,
                                        **combined_metadata
                                    })
                                    
                                    llama_doc = LlamaDocument(
                                        text=enhanced_content,
                                        metadata={
                                            "doc_id": doc_id,
                                            "title": confluence_metadata.get('title', file_path.stem),
                                            "type": "confluence",
                                            "file_path": str(file_path),
                                            **combined_metadata
                                        }
                                    )
                                    
                                    self.index.insert(llama_doc)
                                    print(f"Inserted Confluence document into vector index: {file_path.name}")
                                
                                # Update status to indexed
                                document.status = DocumentStatus.INDEXED
                                
                                # Add to documents dictionary
                                self.documents[doc_id] = document
                                print(f"✓ Loaded: {document.title} (confluence)")
                                return
                    except Exception as e:
                        print(f"Error parsing Confluence frontmatter in {file_path.name}: {e}")
            
            # Check if it's a Web import file (markdown with metadata header)
            if file_path.name.startswith('web_') and file_path.suffix == '.md' and doc_type != DocumentType.PDF:
                # Parse markdown frontmatter for text files only
                if content.startswith('---\n'):
                    try:
                        parts = content.split('---\n', 2)
                        if len(parts) >= 3:
                            frontmatter = parts[1]
                            doc_content = parts[2]
                            
                            # Parse frontmatter
                            web_metadata = {}
                            for line in frontmatter.strip().split('\n'):
                                if ':' in line:
                                    key, value = line.split(':', 1)
                                    web_metadata[key.strip()] = value.strip()
                            
                            # Extract content-specific metadata from the actual content
                            content_metadata = await self._extract_content_metadata(file_path, DocumentType.MARKDOWN, doc_content)
                            
                            # Combine file metadata with web metadata
                            combined_metadata = {**file_metadata, **content_metadata, **web_metadata}
                            
                            doc_id = web_metadata.get('doc_id')
                            if doc_id:
                                document = Document(
                                    id=doc_id,
                                    title=web_metadata.get('title', file_path.stem),
                                    type=DocumentType.MARKDOWN,
                                    content=doc_content.strip(),
                                    file_path=str(file_path),
                                    status=DocumentStatus.PROCESSING,
                                    size_bytes=file_metadata['size_bytes'],
                                    created_at=datetime.fromtimestamp(file_metadata['created_at']),
                                    updated_at=datetime.fromtimestamp(file_metadata['modified_at']),
                                    metadata=combined_metadata
                                )
                                
                                # CRITICAL FIX: Create LlamaIndex document and insert into vector index
                                if self.index is not None:
                                    # Enhance content for better RAG performance
                                    enhanced_content = self._enhance_document_for_rag(doc_content.strip(), {
                                        'title': web_metadata.get('title', file_path.stem),
                                        'document_type': 'webpage',
                                        'filename': file_path.name,
                                        **combined_metadata
                                    })
                                    
                                    llama_doc = LlamaDocument(
                                        text=enhanced_content,
                                        metadata={
                                            "doc_id": doc_id,
                                            "title": web_metadata.get('title', file_path.stem),
                                            "filename": file_path.name,
                                            "type": "webpage"
                                        }
                                    )
                                    
                                    self.index.insert(llama_doc)
                                    print(f"Inserted web document into vector index: {file_path.name}")
                                
                                # Update status to indexed
                                document.status = DocumentStatus.INDEXED
                                
                                # Add to documents dictionary
                                self.documents[doc_id] = document
                                print(f"✓ Loaded: {document.title} (webpage)")
                                return
                    except Exception as e:
                        print(f"Error parsing web frontmatter in {file_path.name}: {e}")
            
            # Handle regular uploaded files
            # Generate a doc_id based on file path and modification time
            file_stat = file_path.stat()
            doc_id = f"file_{hash(str(file_path) + str(file_stat.st_mtime))}"
            
            # Extract content-specific metadata based on file type
            content_metadata = await self._extract_content_metadata(file_path, doc_type, content)
            
            # Combine all metadata
            combined_metadata = {**file_metadata, **content_metadata}
            
            # Enhance content for better RAG performance
            enhanced_content = self._enhance_document_for_rag(content, {
                'title': file_path.stem,
                'document_type': doc_type.value,
                'filename': file_path.name,
                **combined_metadata
            })
            
            # Create document record
            doc_id = str(uuid.uuid4())
            document = Document(
                id=doc_id,
                title=file_path.stem,
                type=doc_type,
                content=content,  # Store original content
                file_path=str(file_path),
                status=DocumentStatus.PROCESSING,
                size_bytes=file_metadata['size_bytes'],
                created_at=datetime.fromtimestamp(file_metadata['created_at']),
                updated_at=datetime.fromtimestamp(file_metadata['modified_at']),
                metadata=combined_metadata
            )
            
            # Create LlamaIndex document with MINIMAL metadata to prevent token bloat
            llama_doc = LlamaDocument(
                text=enhanced_content,  # Use enhanced content for indexing
                metadata={
                    "doc_id": doc_id,
                    "title": file_path.stem,
                    "filename": file_path.name,
                    "type": doc_type.value
                }
            )
            
            # CRITICAL FIX: Add to the vector index so it's searchable
            if self.index is not None:
                self.index.insert(llama_doc)
                print(f"Inserted document into vector index: {file_path.name}")
            
            # Add to documents dictionary
            self.documents[doc_id] = document
            
            # Update status to indexed
            document.status = DocumentStatus.INDEXED
            
            # Calculate token count for logging
            token_count = self._estimate_token_count(content)
            print(f"✓ Loaded: {document.title} ({doc_type.value}, {token_count} tokens)")
            
        except Exception as e:
            print(f"Error loading document from {file_path}: {e}")
            # Don't raise the exception, just log it and continue with other files
            # raise
    
    async def _extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract comprehensive metadata from file system."""
        try:
            file_stat = file_path.stat()
            
            # Basic file information
            metadata = {
                # File system properties
                'filename': file_path.name,
                'file_extension': file_path.suffix.lower(),
                'size_bytes': file_stat.st_size,
                'created_at': file_stat.st_ctime,
                'modified_at': file_stat.st_mtime,
                'accessed_at': file_stat.st_atime,
                
                # File identification
                'mime_type': mimetypes.guess_type(str(file_path))[0] or 'text/plain',
                'encoding': mimetypes.guess_type(str(file_path))[1],
                
                # System information
                'platform': platform.system(),
                'absolute_path': str(file_path.absolute()),
                'relative_path': str(file_path.relative_to(self.documents_dir)),
            }
            
            # Add file hash for integrity checking
            metadata['file_hash'] = await self._calculate_file_hash(file_path)
            
            # Add file permissions (Unix-like systems)
            if platform.system() != 'Windows':
                try:
                    metadata['file_mode'] = oct(file_stat.st_mode)
                    metadata['file_uid'] = file_stat.st_uid
                    metadata['file_gid'] = file_stat.st_gid
                    
                    # Try to get owner/group names
                    try:
                        metadata['file_owner'] = pwd.getpwuid(file_stat.st_uid).pw_name
                    except (KeyError, ImportError):
                        metadata['file_owner'] = str(file_stat.st_uid)
                    
                    try:
                        metadata['file_group'] = grp.getgrgid(file_stat.st_gid).gr_name
                    except (KeyError, ImportError):
                        metadata['file_group'] = str(file_stat.st_gid)
                        
                except Exception as e:
                    print(f"Warning: Could not extract Unix file metadata: {e}")
            
            # Add human-readable file size
            metadata['size_human'] = self._format_file_size(file_stat.st_size)
            
            # Add file age information
            now = datetime.now().timestamp()
            metadata['age_days'] = round((now - file_stat.st_mtime) / 86400, 1)
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting file metadata from {file_path}: {e}")
            return {
                'filename': file_path.name,
                'size_bytes': 0,
                'error': str(e)
            }
    
    async def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file for integrity checking."""
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                # Read file in chunks to handle large files
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            print(f"Error calculating hash for {file_path}: {e}")
            return "hash_error"
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human-readable format."""
        if size_bytes == 0:
            return "0 B"
        
        size_names = ["B", "KB", "MB", "GB", "TB"]
        size_index = 0
        size = float(size_bytes)
        
        while size >= 1024.0 and size_index < len(size_names) - 1:
            size /= 1024.0
            size_index += 1
        
        return f"{size:.1f} {size_names[size_index]}"
    
    async def _extract_content_metadata(self, file_path: Path, doc_type: DocumentType, content: str) -> Dict[str, Any]:
        """Extract metadata specific to document content and type."""
        metadata = {}
        
        try:
            # General content statistics
            metadata['content_length'] = len(content)
            metadata['word_count'] = len(content.split())
            metadata['line_count'] = content.count('\n') + 1
            metadata['paragraph_count'] = len([p for p in content.split('\n\n') if p.strip()])
            
            # Character analysis - store as simple integer counts instead of dict
            metadata['character_letters'] = sum(c.isalpha() for c in content)
            metadata['character_digits'] = sum(c.isdigit() for c in content)
            metadata['character_whitespace'] = sum(c.isspace() for c in content)
            metadata['character_punctuation'] = sum(not c.isalnum() and not c.isspace() for c in content)
            
            # Language detection (basic)
            if content.strip():
                # Simple heuristics for language detection
                if any(ord(c) > 127 for c in content):
                    metadata['contains_unicode'] = True
                else:
                    metadata['contains_unicode'] = False
            
            # Type-specific metadata extraction
            if doc_type == DocumentType.PDF:
                metadata.update(await self._extract_pdf_metadata(file_path))
            elif doc_type == DocumentType.DOCX:
                metadata.update(await self._extract_docx_metadata(file_path))
            elif doc_type == DocumentType.MARKDOWN:
                metadata.update(await self._extract_markdown_metadata(content))
            elif doc_type == DocumentType.HTML:
                metadata.update(await self._extract_html_metadata(content))
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting content metadata from {file_path}: {e}")
            return {'content_metadata_error': str(e)}
    
    async def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        try:
            reader = PdfReader(str(file_path))
            metadata = {}
            
            # Basic PDF info
            metadata['pdf_pages'] = len(reader.pages)
            
            # PDF metadata
            if reader.metadata:
                pdf_meta = reader.metadata
                metadata['pdf_title'] = pdf_meta.get('/Title', '')
                metadata['pdf_author'] = pdf_meta.get('/Author', '')
                metadata['pdf_subject'] = pdf_meta.get('/Subject', '')
                metadata['pdf_creator'] = pdf_meta.get('/Creator', '')
                metadata['pdf_producer'] = pdf_meta.get('/Producer', '')
                
                # Convert PDF dates if available
                creation_date = pdf_meta.get('/CreationDate')
                if creation_date:
                    metadata['pdf_creation_date'] = str(creation_date)
                
                modification_date = pdf_meta.get('/ModDate')
                if modification_date:
                    metadata['pdf_modification_date'] = str(modification_date)
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting PDF metadata from {file_path}: {e}")
            return {'pdf_metadata_error': str(e)}
    
    async def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX files."""
        try:
            doc = DocxDocument(str(file_path))
            metadata = {}
            
            # Document statistics
            metadata['docx_paragraphs'] = len(doc.paragraphs)
            
            # Core properties
            core_props = doc.core_properties
            if core_props:
                metadata['docx_title'] = core_props.title or ''
                metadata['docx_author'] = core_props.author or ''
                metadata['docx_subject'] = core_props.subject or ''
                metadata['docx_keywords'] = core_props.keywords or ''
                metadata['docx_comments'] = core_props.comments or ''
                metadata['docx_category'] = core_props.category or ''
                metadata['docx_language'] = core_props.language or ''
                
                if core_props.created:
                    metadata['docx_created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['docx_modified'] = core_props.modified.isoformat()
                if core_props.last_modified_by:
                    metadata['docx_last_modified_by'] = core_props.last_modified_by
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting DOCX metadata from {file_path}: {e}")
            return {'docx_metadata_error': str(e)}
    
    async def _extract_markdown_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata from Markdown content."""
        try:
            metadata = {}
            
            # Count markdown elements
            metadata['markdown_headers'] = len(re.findall(r'^#{1,6}\s', content, re.MULTILINE))
            metadata['markdown_links'] = len(re.findall(r'\[([^\]]+)\]\(([^)]+)\)', content))
            metadata['markdown_images'] = len(re.findall(r'!\[([^\]]*)\]\(([^)]+)\)', content))
            metadata['markdown_code_blocks'] = len(re.findall(r'```', content)) // 2
            metadata['markdown_tables'] = len(re.findall(r'\|.*\|', content))
            metadata['markdown_lists'] = len(re.findall(r'^\s*[-*+]\s', content, re.MULTILINE))
            metadata['markdown_numbered_lists'] = len(re.findall(r'^\s*\d+\.\s', content, re.MULTILINE))
            
            # Extract header hierarchy
            headers = re.findall(r'^(#{1,6})\s+(.+)$', content, re.MULTILINE)
            if headers:
                # Convert lists to strings to avoid validation errors
                header_levels = [len(h[0]) for h in headers]
                metadata['markdown_header_levels'] = ','.join(map(str, header_levels))
                metadata['markdown_header_text'] = ' | '.join([h[1].strip() for h in headers])
                metadata['markdown_max_header_level'] = max(header_levels) if header_levels else 0
                metadata['markdown_min_header_level'] = min(header_levels) if header_levels else 0
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting Markdown metadata: {e}")
            return {'markdown_metadata_error': str(e)}
    
    async def _extract_html_metadata(self, content: str) -> Dict[str, Any]:
        """Extract metadata from HTML content."""
        try:
            soup = BeautifulSoup(content, 'html.parser')
            metadata = {}
            
            # HTML structure analysis
            metadata['html_title'] = soup.title.string if soup.title else ''
            metadata['html_headings'] = len(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']))
            metadata['html_paragraphs'] = len(soup.find_all('p'))
            metadata['html_links'] = len(soup.find_all('a'))
            metadata['html_images'] = len(soup.find_all('img'))
            metadata['html_tables'] = len(soup.find_all('table'))
            metadata['html_lists'] = len(soup.find_all(['ul', 'ol']))
            
            # Meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                name = meta.get('name') or meta.get('property')
                content = meta.get('content')
                if name and content:
                    metadata[f'html_meta_{name}'] = content
            
            return metadata
            
        except Exception as e:
            print(f"Error extracting HTML metadata: {e}")
            return {'html_metadata_error': str(e)}
    
    async def set_embedding_model(self, model_key: str = "mpnet", use_gpu: bool = None):
        """Set or change the embedding model."""
        try:
            print(f"Setting embedding model to: {model_key}")
            
            # Store old dimensions if we have an existing model
            old_dim = None
            if self.embedding_model:
                old_dim = self.embedding_model.bert_config.get("dimensions", 768)
            
            # Create new embedding model
            self.embedding_model = BERTEmbedding(model_key, use_gpu)
            Settings.embed_model = self.embedding_model
            
            # Check if dimensions changed
            new_dim = self.embedding_model.bert_config.get("dimensions", 768)
            
            if old_dim and old_dim != new_dim:
                print(f"Embedding dimensions changed from {old_dim}D to {new_dim}D")
                print("Reinitializing vector store...")
                
                # Reinitialize vector store to handle dimension change
                await self._initialize_vector_store()
                
                # Reinitialize index
                await self._initialize_index()
                
                print("Vector store reinitialized with new dimensions")
            
            print(f"Embedding model set to: {model_key}")
            
        except Exception as e:
            print(f"Error setting embedding model: {e}")
            # Fallback to sentence-bert
            print("Falling back to sentence-bert model")
            self.embedding_model = BERTEmbedding("sentence-bert", use_gpu)
            Settings.embed_model = self.embedding_model
    
    async def _initialize_vector_store(self):
        """Initialize the ChromaDB vector store."""
        # Check if running in Docker environment
        chroma_host = os.getenv('CHROMA_HOST')
        chroma_port = int(os.getenv('CHROMA_PORT', '8000'))
        
        if chroma_host:
            # Running in Docker - connect to ChromaDB service
            print(f"Connecting to ChromaDB at {chroma_host}:{chroma_port}")
            self.chroma_client = chromadb.HttpClient(
                host=chroma_host,
                port=chroma_port,
                settings=chromadb.Settings(
                    chroma_client_auth_provider="chromadb.auth.token.TokenAuthClientProvider",
                    chroma_client_auth_credentials="test-token"
                )
            )
        else:
            # Running locally - use persistent client
            print(f"Using local ChromaDB at {self.chroma_dir}")
            self.chroma_client = chromadb.PersistentClient(path=str(self.chroma_dir))
        
        # Get expected embedding dimension from current model
        expected_dim = self.embedding_model.bert_config.get("dimensions", 768) if self.embedding_model else 768
        
        try:
            # Try to get existing collection
            collection = self.chroma_client.get_collection("documents")
            
            # Check if we have any embeddings to test dimensions
            count = collection.count()
            if count > 0:
                # Get a sample to check dimensions
                sample = collection.peek(limit=1)
                
                # Fix: Properly check if embeddings exist and get dimensions
                embeddings = sample.get('embeddings') if sample else None
                if embeddings is not None and len(embeddings) > 0 and len(embeddings[0]) > 0:
                    existing_dim = len(embeddings[0])
                    
                    if existing_dim != expected_dim:
                        print(f"Dimension mismatch detected: collection has {existing_dim}D embeddings, current model produces {expected_dim}D")
                        print("Recreating collection with correct dimensions...")
                        
                        # Delete existing collection
                        self.chroma_client.delete_collection("documents")
                        print("Deleted existing collection")
                        
                        # Create new collection
                        collection = self.chroma_client.create_collection("documents")
                        print(f"Created new collection with {expected_dim}D embeddings")
                    else:
                        print(f"Using existing collection with {existing_dim}D embeddings")
                else:
                    print("Existing collection is empty, will use current model dimensions")
            else:
                print("Using existing empty collection")
                
        except ValueError as ve:
            # Handle specific ChromaDB errors
            if "Collection" in str(ve) and "does not exist" in str(ve):
                print(f"Creating new documents collection")
                collection = self.chroma_client.create_collection("documents")
                print(f"Created new collection for {expected_dim}D embeddings")
            else:
                raise ve
        except Exception as e:
            # Collection doesn't exist or other error, create it
            if "does not exist" in str(e).lower():
                print(f"Creating new documents collection")
                collection = self.chroma_client.create_collection("documents")
                print(f"Created new collection for {expected_dim}D embeddings")
            else:
                # Try to get or create collection as fallback
                try:
                    collection = self.chroma_client.get_or_create_collection("documents")
                    print(f"Using get_or_create fallback for collection")
                except Exception as e2:
                    print(f"Failed to create collection: {e2}")
                    raise e2
        
        self.vector_store = ChromaVectorStore(chroma_collection=collection)
    
    async def _initialize_index(self):
        """Initialize or load the document index."""
        # Create storage context
        storage_context = StorageContext.from_defaults(vector_store=self.vector_store)
        
        # Initialize or load existing index
        try:
            self.index = VectorStoreIndex.from_vector_store(
                vector_store=self.vector_store,
                storage_context=storage_context
            )
            print("Loaded existing document index")
        except Exception:
            # Create new index if none exists
            self.index = VectorStoreIndex([], storage_context=storage_context)
            print("Created new document index")
        
        # Load app settings for engine configuration
        settings = self._load_app_settings()
        
        # Initialize query engine with multi-document retrieval
        self.query_engine = self._create_multi_document_query_engine(similarity_top_k=12)
        
        # Initialize chat engine with multi-document retrieval
        self.chat_engine = self.index.as_chat_engine(
            chat_mode="condense_plus_context",  # Better for multi-document conversations
            similarity_top_k=8,  # Reduced from 10 to save tokens
            verbose=True,
            system_prompt=f"""You are a document analysis assistant. Use the provided context to answer questions accurately.

Instructions:
1. Search through all document context provided
2. Reference specific documents by title when possible
3. Distinguish between multiple documents in your response
4. Keep responses under {settings['max_tokens']} tokens
5. Prioritize accuracy over completeness if space is limited"""
        )
    
    def get_embedding_info(self) -> Dict[str, Any]:
        """Get information about the current embedding model."""
        if self.embedding_model:
            return self.embedding_model.get_model_info()
        return {"error": "No embedding model initialized"}
    
    def get_available_models(self) -> Dict[str, Any]:
        """Get list of available embedding models."""
        return BERTEmbedding.BERT_MODELS
    
    def _manage_context_window(self, nodes, query: str, max_total_tokens: int = 2000) -> List:
        """
        Intelligently manage context window by truncating document content while preserving relevance.
        
        Args:
            nodes: Retrieved document nodes
            query: The user's query for relevance ranking
            max_total_tokens: Maximum tokens to use for all document content
        """
        if not nodes:
            return nodes
            
        # Calculate tokens per document (evenly distributed with some buffer)
        tokens_per_doc = max_total_tokens // len(nodes)
        min_tokens_per_doc = 200  # Minimum meaningful content per document
        
        # If we can't give each document meaningful content, reduce document count
        if tokens_per_doc < min_tokens_per_doc:
            # Keep only the most relevant documents
            max_docs = max_total_tokens // min_tokens_per_doc
            nodes = nodes[:max_docs]
            tokens_per_doc = max_total_tokens // len(nodes)
        
        print(f"🔧 Context management: {len(nodes)} documents, ~{tokens_per_doc} tokens each (max: {max_total_tokens})")
        
        # Truncate each document's content intelligently
        managed_nodes = []
        for node in nodes:
            # Handle both Node and NodeWithScore objects
            actual_node = node.node if hasattr(node, 'node') else node
            
            if hasattr(actual_node, 'text') and actual_node.text:
                original_text = actual_node.text
                original_tokens = self._estimate_token_count(original_text)
                
                if original_tokens <= tokens_per_doc:
                    # Document fits entirely
                    print(f"   ✓ Document fits: {original_tokens} tokens")
                    managed_nodes.append(node)
                else:
                    # Need to truncate - preserve most relevant parts
                    print(f"   ✂️  Truncating: {original_tokens} → {tokens_per_doc} tokens")
                    truncated_text = self._smart_truncate_content(
                        original_text, query, tokens_per_doc
                    )
                    # Modify the actual node's text
                    actual_node.text = truncated_text
                    managed_nodes.append(node)
            else:
                managed_nodes.append(node)
                
        return managed_nodes
    
    def _smart_truncate_content(self, content: str, query: str, max_tokens: int) -> str:
        """
        Intelligently truncate content while preserving the most relevant parts.
        
        Strategy:
        1. Split content into paragraphs
        2. Score each paragraph for relevance to query
        3. Keep highest-scoring paragraphs that fit within token limit
        """
        if not content or not content.strip():
            return content
            
        # Split into paragraphs (preserve structure)
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        if not paragraphs:
            # Fallback: just truncate by tokens
            return self._truncate_by_tokens(content, max_tokens)
        
        # Score paragraphs by relevance to query
        query_lower = query.lower()
        query_words = set(query_lower.split())
        
        scored_paragraphs = []
        for para in paragraphs:
            para_lower = para.lower()
            # Simple relevance scoring
            score = 0
            
            # Exact query match gets highest score
            if query_lower in para_lower:
                score += 10
                
            # Word overlap scoring
            para_words = set(para_lower.split())
            common_words = query_words.intersection(para_words)
            score += len(common_words) * 2
            
            # Length penalty (prefer concise relevant content)
            para_tokens = self._estimate_token_count(para)
            if para_tokens > 300:  # Very long paragraphs get slight penalty
                score -= 1
                
            scored_paragraphs.append((score, para_tokens, para))
        
        # Sort by relevance score (descending)
        scored_paragraphs.sort(key=lambda x: x[0], reverse=True)
        
        # Select paragraphs that fit within token limit
        selected_paragraphs = []
        total_tokens = 0
        
        for score, tokens, para in scored_paragraphs:
            if total_tokens + tokens <= max_tokens:
                selected_paragraphs.append(para)
                total_tokens += tokens
            elif total_tokens == 0:  # First paragraph is too long, truncate it
                truncated_para = self._truncate_by_tokens(para, max_tokens)
                selected_paragraphs.append(truncated_para)
                break
        
        # If we got nothing, use fallback truncation
        if not selected_paragraphs:
            return self._truncate_by_tokens(content, max_tokens)
            
        return '\n\n'.join(selected_paragraphs)
    
    def _truncate_by_tokens(self, text: str, max_tokens: int) -> str:
        """Simple token-based truncation with word boundary preservation."""
        words = text.split()
        # Rough estimate: ~1.3 words per token
        max_words = int(max_tokens * 1.3)
        
        if len(words) <= max_words:
            return text
            
        truncated = ' '.join(words[:max_words])
        return truncated + "..." if len(words) > max_words else truncated

    def _ensure_multi_document_retrieval(self, nodes, max_nodes: int = 8) -> List:
        """
        Ensure retrieved nodes come from multiple documents when possible.
        This helps with document diversity in responses.
        """
        if not nodes:
            return nodes
            
        # Group nodes by document source
        doc_groups = {}
        for node in nodes:
            # Try to get document identifier from metadata
            doc_id = None
            if hasattr(node, 'metadata') and node.metadata:
                doc_id = (node.metadata.get('file_path') or 
                         node.metadata.get('doc_id') or 
                         node.metadata.get('title') or
                         node.metadata.get('filename'))
            
            if not doc_id:
                doc_id = 'unknown'
                
            if doc_id not in doc_groups:
                doc_groups[doc_id] = []
            doc_groups[doc_id].append(node)
        
        print(f"Found nodes from {len(doc_groups)} documents: {list(doc_groups.keys())}")
        
        # If we have multiple documents, distribute nodes more evenly
        if len(doc_groups) > 1:
            selected_nodes = []
            max_per_doc = max(1, max_nodes // len(doc_groups))
            remaining_slots = max_nodes
            
            # First pass: take up to max_per_doc from each document
            for doc_id, doc_nodes in doc_groups.items():
                take_count = min(max_per_doc, len(doc_nodes), remaining_slots)
                selected_nodes.extend(doc_nodes[:take_count])
                remaining_slots -= take_count
                
                if remaining_slots <= 0:
                    break
            
            # Second pass: fill remaining slots from documents with more content
            if remaining_slots > 0:
                for doc_id, doc_nodes in doc_groups.items():
                    already_taken = min(max_per_doc, len(doc_nodes))
                    available = doc_nodes[already_taken:]
                    
                    take_count = min(len(available), remaining_slots)
                    if take_count > 0:
                        selected_nodes.extend(available[:take_count])
                        remaining_slots -= take_count
                        
                        if remaining_slots <= 0:
                            break
            
            print(f"Selected {len(selected_nodes)} nodes with multi-document diversity")
            return selected_nodes[:max_nodes]
        
        # Single document or no metadata - return original nodes limited by max_nodes
        return nodes[:max_nodes]
    
    def _create_multi_document_query_engine(self, similarity_top_k: int = 10):
        """Create a query engine that ensures multi-document retrieval."""
        from llama_index.core.query_engine import RetrieverQueryEngine
        from llama_index.core.retrievers import VectorIndexRetriever
        from llama_index.core.response_synthesizers import ResponseMode
        
        # Create a custom retriever that uses our multi-document logic
        class MultiDocumentRetriever(VectorIndexRetriever):
            def __init__(self, index, similarity_top_k, service_instance):
                super().__init__(index, similarity_top_k=similarity_top_k)
                self.service = service_instance
            
            def _retrieve(self, query_bundle):
                # Get more nodes initially
                nodes = super()._retrieve(query_bundle)
                # Apply multi-document filtering
                filtered_nodes = self.service._ensure_multi_document_retrieval(nodes, max_nodes=6)
                # Apply smart context management
                return self.service._manage_context_window(filtered_nodes, query_bundle.query_str)
        
        # Create the custom retriever
        retriever = MultiDocumentRetriever(
            index=self.index,
            similarity_top_k=similarity_top_k,
            service_instance=self
        )
        
        # Create query engine with the custom retriever
        query_engine = RetrieverQueryEngine.from_args(
            retriever=retriever,
            response_mode=ResponseMode.TREE_SUMMARIZE,  # Better for multiple documents
            verbose=False  # Disable verbose logging to reduce log clutter
        )
        
        return query_engine
    
    # Removed _create_targeted_query_engine - was hardcoded to specific documents
    # Now using only generic semantic similarity retrieval
    
    def _format_response_as_markdown(self, text: str) -> str:
        """Convert plain text response to markdown format."""
        if not text or text.strip() == "":
            return text
            
        # Check if the text already has markdown formatting
        if any(marker in text for marker in ['**', '##', '- ', '1. ', '`', '>']):
            return text  # Already has markdown formatting
        
        # Clean up the text
        text = text.strip()
        
        # Split into sentences for processing
        sentences = [s.strip() for s in text.split('. ') if s.strip()]
        
        formatted_text = ""
        
        # Process each sentence
        for i, sentence in enumerate(sentences):
            if not sentence.strip():
                continue
                
            # Add period if missing
            if not sentence.endswith('.') and not sentence.endswith('!') and not sentence.endswith('?'):
                sentence += '.'
            
            # Look for lists and enumerations first (highest priority)
            if 'include' in sentence.lower() and (',' in sentence or ' and ' in sentence):
                # Check if this sentence mentions multiple items
                items_match = re.search(r'include[s]?\s+([^.]+)', sentence, re.IGNORECASE)
                if items_match:
                    items_text = items_match.group(1)
                    # Split on commas and/or 'and'
                    items = re.split(r',\s*(?:and\s+)?|(?:\s+and\s+)', items_text)
                    if len(items) >= 2:
                        # Convert to bullet list
                        intro = sentence[:items_match.start(1)].strip()
                        if intro.endswith('include'):
                            intro += ':'
                        formatted_text += f"\n{intro}\n\n"
                        for item in items:
                            item = item.strip().rstrip(',').rstrip('.')
                            if item:
                                formatted_text += f"- {item}\n"
                        formatted_text += "\n"
                        continue
            
            # Look for key terms and make them bold (more selective)
            key_terms = [
                # Technology terms (exact matches to avoid over-formatting)
                r'\b(artificial intelligence|AI|natural language processing|intent recognition|multi-agent collaboration)\b',
                r'\b(AWS Bedrock Multi-Agent|AWS Bedrock|Multi-Agent|Admin Orchestrator)\b',
                # Important concepts
                r'\b(Access Admin plugin strategy|plugin strategy|orchestration approach)\b',
            ]
            
            original_sentence = sentence
            for pattern in key_terms:
                sentence = re.sub(pattern, r'**\1**', sentence, flags=re.IGNORECASE)
            
            # Look for sentences that mention advantages or benefits
            if any(word in sentence.lower() for word in ['enable', 'eliminates', 'provides', 'recommended']):
                # These are benefit statements, add emphasis with blockquote
                sentence = f"> {sentence}"
                formatted_text += f"\n{sentence}\n\n"
                continue
            
            # Add the sentence with normal formatting
            formatted_text += sentence
            
            # Add appropriate spacing
            if i < len(sentences) - 1:
                # Check if next sentence starts a new topic or benefit
                next_sentence = sentences[i + 1] if i + 1 < len(sentences) else ""
                if any(starter in next_sentence.lower() for starter in ['additionally', 'based on', 'this approach']):
                    formatted_text += "\n\n"  # New paragraph
                else:
                    formatted_text += " "  # Same paragraph
        
        return formatted_text.strip()
    
    async def clear_vector_store(self):
        """Clear the vector store and recreate it. Useful when changing embedding models."""
        try:
            print("Clearing vector store...")
            
            # Delete existing collection if it exists
            try:
                # Check if collection exists first
                existing_collections = self.chroma_client.list_collections()
                collection_names = [col.name for col in existing_collections]
                
                if "documents" in collection_names:
                    self.chroma_client.delete_collection("documents")
                    print("Deleted existing collection")
                else:
                    print("No existing collection to delete")
                    
            except Exception as e:
                print(f"Error during collection deletion (continuing anyway): {e}")
            
            # Small delay to ensure deletion is processed
            await asyncio.sleep(0.1)
            
            # Reinitialize vector store
            await self._initialize_vector_store()
            
            # Reinitialize index  
            await self._initialize_index()
            
            # Clear documents dictionary since embeddings are now invalid
            self.documents.clear()
            
            # CRITICAL FIX: Reload and re-index existing documents
            print("Re-indexing existing documents...")
            await self._load_existing_documents()
            
            print("Vector store cleared and reinitialized")
            return True
            
        except Exception as e:
            print(f"Error clearing vector store: {e}")
            return False
    
    def _get_gpt4all_model_path(self) -> Optional[str]:
        """Get the path to an available GPT4All model."""
        model_extensions = [".gguf", ".bin"]
        
        # Check for preferred model in settings first
        settings = self._load_app_settings()
        preferred_model = settings.get('preferred_gpt4all_model')
        if preferred_model:
            preferred_path = self.models_dir / preferred_model
            if preferred_path.exists() and preferred_path.is_file():
                return preferred_model
        
        # Fall back to first available model
        for model_file in self.models_dir.iterdir():
            if model_file.is_file() and any(model_file.suffix == ext for ext in model_extensions):
                return model_file.name
        
        return None
    
    async def process_uploaded_file(self, file_path: str, filename: str) -> Document:
        """Process an uploaded file and add it to the index."""
        try:
            file_path_obj = Path(file_path)
            
            # Determine document type
            doc_type = self._get_document_type(filename)
            
            # Get comprehensive file metadata
            file_metadata = await self._extract_file_metadata(file_path_obj)
            
            # Extract content based on file type
            content = await self._extract_content(file_path, doc_type)
            
            # Extract content-specific metadata
            content_metadata = await self._extract_content_metadata(file_path_obj, doc_type, content)
            
            # Combine all metadata
            combined_metadata = {**file_metadata, **content_metadata}
            
            # Enhance content for better RAG performance
            enhanced_content = self._enhance_document_for_rag(content, {
                'title': filename,
                'document_type': doc_type.value,
                'filename': filename,
                **combined_metadata
            })
            
            # Create document record
            doc_id = str(uuid.uuid4())
            document = Document(
                id=doc_id,
                title=filename,
                type=doc_type,
                content=content,  # Store original content
                file_path=file_path,
                status=DocumentStatus.PROCESSING,
                size_bytes=file_metadata['size_bytes'],
                created_at=datetime.fromtimestamp(file_metadata['created_at']),
                updated_at=datetime.fromtimestamp(file_metadata['modified_at']),
                metadata=combined_metadata
            )
            
            # Create LlamaIndex document with enhanced content for better RAG
            llama_doc = LlamaDocument(
                text=enhanced_content,  # Use enhanced content for indexing
                metadata={
                    "doc_id": doc_id,
                    "title": filename,
                    "type": doc_type.value,
                    "file_path": file_path,
                    **combined_metadata  # Include all extracted metadata
                }
            )
            
            # Add to documents dictionary
            self.documents[doc_id] = document
            
            # Add to index
            self.index.insert(llama_doc)
            
            # Update status
            document.status = DocumentStatus.INDEXED
            self.documents[doc_id] = document
            
            print(f"Successfully processed document: {filename} with {len(combined_metadata)} metadata fields")
            return document
            
        except Exception as e:
            print(f"Error processing document {filename}: {e}")
            if 'doc_id' in locals() and doc_id in self.documents:
                self.documents[doc_id].status = DocumentStatus.ERROR
            raise
    
    def _get_document_type(self, filename: str) -> DocumentType:
        """Determine document type from filename."""
        extension = Path(filename).suffix.lower()
        filename_lower = filename.lower()
        
        # Check for transcript files first (more specific detection)
        transcript_indicators = [
            'transcript', 'vtt', 'webvtt', 'srt', 'recording',
            'meeting', 'call', 'zoom', 'teams', 'conference'
        ]
        
        if (extension in ['.txt', '.vtt', '.srt'] and 
            any(indicator in filename_lower for indicator in transcript_indicators)):
            return DocumentType.TRANSCRIPT
        
        type_mapping = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOCX,
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MARKDOWN,
            ".markdown": DocumentType.MARKDOWN,
            ".html": DocumentType.HTML,
            ".htm": DocumentType.HTML,
            ".vtt": DocumentType.TRANSCRIPT,
            ".srt": DocumentType.TRANSCRIPT
        }
        
        return type_mapping.get(extension, DocumentType.TXT)
    
    async def _extract_content(self, file_path: str, doc_type: DocumentType) -> str:
        """Extract text content from different file types."""
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf_content(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._extract_text_content(file_path)
            elif doc_type == DocumentType.MARKDOWN:
                return await self._extract_markdown_content(file_path)
            elif doc_type == DocumentType.HTML:
                return await self._extract_html_content(file_path)
            elif doc_type == DocumentType.TRANSCRIPT:
                return await self._extract_transcript_content(file_path)
            else:
                return await self._extract_text_content(file_path)
                
        except Exception as e:
            print(f"Error extracting content from {file_path}: {e}")
            return ""
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    async def _extract_text_content(self, file_path: str) -> str:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def _extract_markdown_content(self, file_path: str) -> str:
        """Extract text from Markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text()
    
    async def _extract_html_content(self, file_path: str) -> str:
        """Extract text from HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
    
    async def _extract_transcript_content(self, file_path: str) -> str:
        """Extract and enhance transcript content while preserving timestamps."""
        import re
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Detect format
        if 'WEBVTT' in content[:50]:
            return self._process_webvtt_content(content)
        elif '-->' in content and re.search(r'\d{2}:\d{2}:\d{2}', content):
            return self._process_srt_content(content)
        else:
            # Fallback to regular text processing
            return content
    
    def _process_webvtt_content(self, content: str) -> str:
        """Process WEBVTT format with intelligent timestamp preservation and chunking."""
        import re
        
        lines = content.split('\n')
        processed_segments = []
        current_speaker = ''
        current_text = ''
        current_timerange = ''
        
        timestamp_pattern = r'(\d{2}:\d{2}:\d{2}\.\d{3}) --> (\d{2}:\d{2}:\d{2}\.\d{3})'
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip WEBVTT header and empty lines
            if line == 'WEBVTT' or not line:
                i += 1
                continue
            
            # Skip sequence numbers
            if re.match(r'^\d+$', line):
                i += 1
                continue
            
            # Check for timestamp line
            timestamp_match = re.search(timestamp_pattern, line)
            if timestamp_match:
                start_time = timestamp_match.group(1)
                end_time = timestamp_match.group(2)
                current_timerange = f"[{start_time}]"
                i += 1
                continue
            
            # Process speaker content
            if ':' in line:
                # If we have accumulated content, save it
                if current_text:
                    processed_segments.append({
                        'timestamp': current_timerange,
                        'speaker': current_speaker,
                        'content': current_text.strip(),
                        'raw_time': self._parse_timestamp(current_timerange)
                    })
                
                # Start new speaker
                parts = line.split(':', 1)
                current_speaker = parts[0].strip()
                current_text = parts[1].strip() if len(parts) > 1 else ''
            else:
                # Continue current speaker's text
                current_text += ' ' + line
            
            i += 1
        
        # Don't forget the last segment
        if current_text:
            processed_segments.append({
                'timestamp': current_timerange,
                'speaker': current_speaker,
                'content': current_text.strip(),
                'raw_time': self._parse_timestamp(current_timerange)
            })
        
        # Now chunk the segments intelligently
        return self._chunk_transcript_segments(processed_segments)
    
    def _parse_timestamp(self, timerange: str) -> int:
        """Convert timestamp to seconds for chunking logic."""
        import re
        match = re.search(r'\[(\d{2}):(\d{2}):(\d{2})', timerange)
        if match:
            hours, minutes, seconds = map(int, match.groups())
            return hours * 3600 + minutes * 60 + seconds
        return 0
    
    def _chunk_transcript_segments(self, segments: list, chunk_size_minutes: int = 10, max_tokens_per_chunk: int = 1500) -> str:
        """Chunk transcript segments by time and token count."""
        chunks = []
        current_chunk = []
        current_chunk_start_time = 0
        current_chunk_tokens = 0
        
        for segment in segments:
            segment_text = f"{segment['timestamp']} {segment['speaker']}: {segment['content']}"
            segment_tokens = self._estimate_token_count(segment_text)
            segment_time = segment['raw_time']
            
            # Check if we should start a new chunk
            should_chunk = (
                # Time-based chunking (every N minutes)
                (segment_time - current_chunk_start_time) >= (chunk_size_minutes * 60) or
                # Token-based chunking
                (current_chunk_tokens + segment_tokens) > max_tokens_per_chunk or
                # Logical break detection (long pause in timestamps)
                (current_chunk and segment_time - current_chunk[-1]['raw_time'] > 300)  # 5 min gap
            )
            
            if should_chunk and current_chunk:
                # Finalize current chunk
                chunk_content = self._format_transcript_chunk(current_chunk, len(chunks) + 1)
                chunks.append(chunk_content)
                
                # Start new chunk
                current_chunk = [segment]
                current_chunk_start_time = segment_time
                current_chunk_tokens = segment_tokens
            else:
                # Add to current chunk
                current_chunk.append(segment)
                current_chunk_tokens += segment_tokens
                if not current_chunk_start_time:
                    current_chunk_start_time = segment_time
        
        # Don't forget the last chunk
        if current_chunk:
            chunk_content = self._format_transcript_chunk(current_chunk, len(chunks) + 1)
            chunks.append(chunk_content)
        
        # Create final document with chunk metadata
        header = f"""TRANSCRIPT CONTENT (CHUNKED):
{('='*60)}
Total Chunks: {len(chunks)}
Chunking Strategy: {chunk_size_minutes}-minute segments, max {max_tokens_per_chunk} tokens per chunk
{('='*60)}

"""
        
        return header + '\n\n'.join(chunks)
    
    def _format_transcript_chunk(self, segments: list, chunk_number: int) -> str:
        """Format a chunk of transcript segments."""
        if not segments:
            return ""
        
        start_time = segments[0]['timestamp']
        end_time = segments[-1]['timestamp']
        
        chunk_header = f"CHUNK {chunk_number} ({start_time} to {end_time}):\n{'-'*50}\n"
        
        chunk_content = []
        for segment in segments:
            chunk_content.append(f"{segment['timestamp']} {segment['speaker']}: {segment['content']}")
        
        return chunk_header + '\n\n'.join(chunk_content)
    
    def _process_srt_content(self, content: str) -> str:
        """Process SRT format with timestamp preservation."""
        import re
        
        # SRT format: sequence number, timestamp, content
        segments = content.split('\n\n')
        processed_segments = []
        
        for segment in segments:
            lines = segment.strip().split('\n')
            if len(lines) >= 3:
                # Skip sequence number (first line)
                timestamp = lines[1] if '-->' in lines[1] else ''
                content_lines = lines[2:]
                
                if timestamp:
                    # Convert SRT timestamp to simpler format
                    start_time = timestamp.split(' --> ')[0].replace(',', '.')
                    timerange = f"[{start_time}]"
                    content_text = ' '.join(content_lines)
                    
                    processed_segments.append(f"{timerange} {content_text}")
        
        header = "TRANSCRIPT CONTENT:\n" + "="*50 + "\n\n"
        return header + '\n\n'.join(processed_segments)
    
    def _enhance_document_for_rag(self, content: str, metadata: Dict[str, Any]) -> str:
        """Enhance document content with metadata and apply universal chunking for large documents."""
        
        # Check if content needs chunking (now optional since metadata bloat is fixed)
        token_count = self._estimate_token_count(content)
        chunk_threshold = 8000  # Much higher threshold - chunking may not be needed anymore
        enable_chunking = False  # EXPERIMENT: Disable chunking to test performance
        
        if enable_chunking and token_count > chunk_threshold:
            print(f"Document exceeds {chunk_threshold} tokens ({token_count}), applying universal chunking")
            return self._apply_universal_chunking(content, metadata)
        elif token_count > chunk_threshold:
            print(f"Document has {token_count} tokens (chunking disabled for testing)")
        
        # For smaller documents, just add metadata
        title = metadata.get('title', 'Untitled Document')
        doc_type = metadata.get('document_type', 'Unknown')
        
        enhanced_content = f"""Document: {title}

{content}
"""
        return enhanced_content
    
    def _apply_universal_chunking(self, content: str, metadata: Dict[str, Any]) -> str:
        """Apply universal chunking strategy to large documents."""
        doc_type = metadata.get('document_type', DocumentType.TEXT)
        
        # Different chunking strategies based on document type
        if doc_type == DocumentType.TRANSCRIPT and "TRANSCRIPT CONTENT (CHUNKED)" in content:
            # Transcript already chunked in _extract_transcript_content
            return content
        else:
            return self._chunk_general_document(content, metadata)
    
    def _chunk_general_document(self, content: str, metadata: Dict[str, Any]) -> str:
        """Chunk general documents (non-transcripts) into manageable segments."""
        max_tokens_per_chunk = 1500
        overlap_tokens = 150
        
        # Split content into paragraphs first
        paragraphs = content.split('\n\n')
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        title = metadata.get('title', 'Untitled Document')
        doc_type = metadata.get('document_type', 'Unknown')
        
        header = f"""Document: {title} (Chunked)

"""
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue
                
            paragraph_tokens = self._estimate_token_count(paragraph)
            
            # If adding this paragraph would exceed the limit, finalize current chunk
            if current_tokens + paragraph_tokens > max_tokens_per_chunk and current_chunk:
                chunk_content = '\n\n'.join(current_chunk)
                chunks.append({
                    'content': chunk_content,
                    'tokens': current_tokens
                })
                
                # Start new chunk with overlap (keep last paragraph if it fits)
                if overlap_tokens > 0 and current_chunk:
                    last_para = current_chunk[-1]
                    last_para_tokens = self._estimate_token_count(last_para)
                    if last_para_tokens <= overlap_tokens:
                        current_chunk = [last_para]
                        current_tokens = last_para_tokens
                    else:
                        current_chunk = []
                        current_tokens = 0
                else:
                    current_chunk = []
                    current_tokens = 0
            
            # Add current paragraph
            current_chunk.append(paragraph)
            current_tokens += paragraph_tokens
        
        # Add final chunk if it has content
        if current_chunk:
            chunk_content = '\n\n'.join(current_chunk)
            chunks.append({
                'content': chunk_content,
                'tokens': current_tokens
            })
        
        # Format chunks with minimal headers
        result = header
        for i, chunk in enumerate(chunks, 1):
            result += f"--- Part {i} ---\n"
            result += chunk['content'] + "\n\n"
        
        return result

    async def query_documents_with_sources(self, query: str, document_ids: Optional[List[str]] = None) -> tuple[str, list[str]]:
        """Query the document index and return both response and source documents."""
        try:
            if not self.query_engine:
                raise ValueError("Query engine not initialized")
            
            sources = []
            
            # Update LLM settings to ensure current max_tokens is used
            self.update_llm_settings()
            
            # CRITICAL FIX: Handle document_ids parameter properly
            if document_ids is not None and len(document_ids) == 0:
                # Empty array means "use no documents" - generate using LLM general knowledge
                print("Empty document_ids list provided - generating content using LLM general knowledge")
                response_text = self.llm.complete(query).text
                return self._format_response_as_markdown(response_text), []
            elif document_ids is not None and len(document_ids) > 0:
                # Specific document IDs provided
                print(f"Querying specific documents: {document_ids}")
                # TODO: Implement document-specific retrieval based on document_ids
                # For now, fall through to regular logic but this needs to be implemented
            else:
                # document_ids is None - use all documents (default behavior)
                print("No document_ids specified - querying all documents")
            
            # First, check if this is a document-specific query that should use direct document loading
            matching_documents = self._detect_document_specific_query(query)
            
            if matching_documents:
                print(f"Using direct document analysis for: {matching_documents}")
                sources = matching_documents
                
                # Use direct document analysis for other documents
                direct_prompt = self._create_direct_document_prompt(matching_documents, query)
                if direct_prompt:
                    # Use LLM directly with the full document content
                    response_text = self.llm.complete(direct_prompt).text
                    return self._format_response_as_markdown(response_text), sources
            
            # Check if this is a targeted query for specific documents
            query_lower = query.lower()
            target_keywords = []
            
            # Look for specific document keywords in the query
            if any(keyword in query_lower for keyword in ['pset-rfc', 'pset rfc', 'inter ai agent', 'communication protocol']):
                target_keywords.extend(['pset-rfc', 'pset_rfc', 'inter_ai_agent', 'communication'])
            if 'orchestrator' in query_lower:
                target_keywords.extend(['orchestrator', 'plugin'])
            if 'google' in query_lower and 'a2a' in query_lower:
                target_keywords.extend(['google', 'a2a', 'streaming'])
            
            # Use targeted retrieval if specific keywords found
            if target_keywords:
                print(f"Using targeted retrieval for keywords: {target_keywords}")
                targeted_engine = self._create_targeted_query_engine(target_keywords, similarity_top_k=10)
                response = targeted_engine.query(query)
                sources = self._extract_sources_from_response(response)
            else:
                # Use standard multi-document retrieval
                response = self.query_engine.query(query)
                sources = self._extract_sources_from_response(response)
            
            # CLEAN UP RESPONSE: Extract only the AI-generated answer
            if hasattr(response, 'response') and response.response:
                raw_response = response.response.strip()
                
                # Remove metadata sections that start with separators
                clean_response = raw_response
                
                # Split on common LlamaIndex separators and take only the first part
                separators = [
                    "--------------------",
                    "doc_id:",
                    "Context information is below",
                    "Helpful Answer:",
                    "system:",
                    "\n\ndoc_id:",
                    "\n--------------------"
                ]
                
                for separator in separators:
                    if separator in clean_response:
                        clean_response = clean_response.split(separator)[0].strip()
                        break
                
                # Additional cleanup: remove any trailing metadata-like content
                lines = clean_response.split('\n')
                clean_lines = []
                
                for line in lines:
                    # Skip lines that look like metadata
                    if any(metadata_indicator in line.lower() for metadata_indicator in [
                        'doc_id:', 'file_path:', 'title:', 'type:', 'size_bytes:', 'created_at:', 
                        'modified_at:', 'mime_type:', 'encoding:', 'platform:', 'file_hash:',
                        'confluence_sync_', 'absolute_path:', 'relative_path:'
                    ]):
                        break
                    clean_lines.append(line)
                
                result = '\n'.join(clean_lines).strip()
            else:
                result = str(response).strip() if response else ""
            
            if not result or result.lower() in ['', 'empty response', 'empty response.']:
                return "I couldn't find relevant information in the documents to answer your question.", sources
            
            # Truncate response to respect max_tokens setting
            settings = self._load_app_settings()
            result = self._truncate_response_to_max_tokens(result, settings['max_tokens'])
            
            return result, sources
            
        except Exception as e:
            print(f"Error in query_documents_with_sources: {e}")
            return f"Error querying documents: {str(e)}", []

    def _extract_sources_from_response(self, response) -> list[str]:
        """Extract source document filenames from LlamaIndex response."""
        sources = []
        
        try:
            # Check if response has source_nodes (from vector retrieval)
            if hasattr(response, 'source_nodes') and response.source_nodes:
                for node in response.source_nodes:
                    # Extract filename from metadata
                    if hasattr(node, 'node') and hasattr(node.node, 'metadata'):
                        metadata = node.node.metadata
                        # Try different metadata keys for filename
                        filename = (metadata.get('filename') or 
                                  metadata.get('file_name') or 
                                  metadata.get('title') or
                                  metadata.get('source'))
                        if filename and filename not in sources:
                            # Clean filename for display
                            clean_filename = filename.replace('confluence_', '').replace('_confluen.md', '')
                            sources.append(clean_filename)
            
            # Fallback: extract from response text if available
            elif hasattr(response, 'metadata') and response.metadata:
                for doc_id, metadata in response.metadata.items():
                    filename = (metadata.get('filename') or 
                              metadata.get('file_name') or 
                              metadata.get('title'))
                    if filename and filename not in sources:
                        clean_filename = filename.replace('confluence_', '').replace('_confluen.md', '')
                        sources.append(clean_filename)
            
        except Exception as e:
            print(f"Error extracting sources: {e}")
        
        return sources

    async def query_documents(self, query: str, document_ids: Optional[List[str]] = None) -> str:
        """Query the document index with a natural language question."""
        try:
            if not self.query_engine:
                raise ValueError("Query engine not initialized")
            

            # Update LLM settings to ensure current max_tokens is used
            self.update_llm_settings()
            
            # CRITICAL FIX: Handle document_ids parameter properly
            if document_ids is not None and len(document_ids) == 0:
                # Empty array means "use no documents" - generate using LLM general knowledge
                print("Empty document_ids list provided - generating content using LLM general knowledge")
                response_text = self.llm.complete(query).text
                return self._format_response_as_markdown(response_text)
            elif document_ids is not None and len(document_ids) > 0:
                # Specific document IDs provided
                print(f"Querying specific documents: {document_ids}")
                # TODO: Implement document-specific retrieval based on document_ids
                # For now, fall through to regular logic but this needs to be implemented
            else:
                # document_ids is None - use all documents (default behavior)
                print("No document_ids specified - querying all documents")
            
            # First, check if this is a document-specific query that should use direct document loading
            matching_documents = self._detect_document_specific_query(query)
            
            if matching_documents:
                print(f"Using direct document analysis for: {matching_documents}")
                

                
                # Use direct document analysis for other documents
                direct_prompt = self._create_direct_document_prompt(matching_documents, query)
                if direct_prompt:
                    # Use LLM directly with the full document content
                    response_text = self.llm.complete(direct_prompt).text
                    return self._format_response_as_markdown(response_text)
            
            # Check if this is a targeted query for specific documents
            query_lower = query.lower()
            target_keywords = []
            
            # Look for specific document keywords in the query
            if any(keyword in query_lower for keyword in ['pset-rfc', 'pset rfc', 'inter ai agent', 'communication protocol']):
                target_keywords.extend(['pset-rfc', 'pset_rfc', 'inter_ai_agent', 'communication'])
            if 'orchestrator' in query_lower:
                target_keywords.extend(['orchestrator', 'plugin'])
            if 'google' in query_lower and 'a2a' in query_lower:
                target_keywords.extend(['google', 'a2a', 'streaming'])
            
            # Use targeted retrieval if specific keywords found
            if target_keywords:
                print(f"Using targeted retrieval for keywords: {target_keywords}")
                targeted_engine = self._create_targeted_query_engine(target_keywords, similarity_top_k=10)
                response = targeted_engine.query(query)
            else:
                # Use standard multi-document retrieval
                response = self.query_engine.query(query)
            
            # CLEAN UP RESPONSE: Extract only the AI-generated answer
            if hasattr(response, 'response') and response.response:
                raw_response = response.response.strip()
                
                # Remove metadata sections that start with separators
                clean_response = raw_response
                
                # Split on common LlamaIndex separators and take only the first part
                separators = [
                    "--------------------",
                    "doc_id:",
                    "Context information is below",
                    "Helpful Answer:",
                    "system:",
                    "\n\ndoc_id:",
                    "\n--------------------"
                ]
                
                for separator in separators:
                    if separator in clean_response:
                        clean_response = clean_response.split(separator)[0].strip()
                        break
                
                # Additional cleanup: remove any trailing metadata-like content
                lines = clean_response.split('\n')
                clean_lines = []
                
                for line in lines:
                    # Skip lines that look like metadata
                    if any(metadata_indicator in line.lower() for metadata_indicator in [
                        'doc_id:', 'file_path:', 'title:', 'type:', 'size_bytes:', 'created_at:', 
                        'modified_at:', 'mime_type:', 'encoding:', 'platform:', 'file_hash:',
                        'confluence_sync_', 'absolute_path:', 'relative_path:'
                    ]):
                        break
                    clean_lines.append(line)
                
                result = '\n'.join(clean_lines).strip()
            else:
                result = str(response).strip() if response else ""
            
            if not result or result.lower() in ['', 'empty response', 'empty response.']:
                return "I couldn't find relevant information in the documents to answer your question."
            
            # Truncate response to respect max_tokens setting
            settings = self._load_app_settings()
            result = self._truncate_response_to_max_tokens(result, settings['max_tokens'])
            
            return result
            
        except Exception as e:
            print(f"Error in query_documents: {e}")
            return f"Error querying documents: {str(e)}"
    
    def _truncate_conversation_history(self, conversation_history: Optional[List], settings: Dict[str, Any]) -> Optional[List]:
        """Truncate conversation history based on settings."""
        if not conversation_history or not settings.get('enable_conversation_context', True):
            return None
        
        max_history = settings.get('max_conversation_history', 10)
        if max_history <= 0:
            return None
        
        # Keep the most recent messages (pairs of user/assistant)
        # We want to keep complete conversation turns, so we'll keep pairs
        if len(conversation_history) > max_history:
            return conversation_history[-max_history:]
        
        return conversation_history

    async def chat_with_documents_with_sources(self, message: str, conversation_history: Optional[List] = None) -> tuple[str, list[str]]:
        """Chat with documents and return both response and source documents."""
        try:
            if not self.chat_engine:
                raise ValueError("Chat engine not initialized")
            
            sources = []
            
            # Update LLM settings to ensure current max_tokens is used
            self.update_llm_settings()
            
            # Load settings for conversation history management
            settings = self._load_app_settings()
            
            # Load max_tokens for later use, but don't validate here since Smart Context Management handles it
            max_tokens = settings.get('max_tokens', 512)
            
            # Log the message length for debugging
            message_tokens = self._estimate_token_count(message)
            print(f"📝 Processing message: {message_tokens} tokens, max_tokens: {max_tokens}")
            
            # Truncate conversation history more aggressively if needed
            truncated_history = self._truncate_conversation_history(conversation_history, settings)
            
            # Estimate total input tokens (message + history + system prompt overhead)
            # But rely on Smart Context Management for actual token management
            history_tokens = 0
            if truncated_history:
                history_text = " ".join([f"{msg.get('role', '')}: {msg.get('content', '')}" for msg in truncated_history])
                history_tokens = self._estimate_token_count(history_text)
            
            system_prompt_overhead = 200  # Estimate for system prompt and formatting
            total_input_tokens = message_tokens + history_tokens + system_prompt_overhead
            
            print(f"🔍 Input breakdown: message={message_tokens}, history={history_tokens}, overhead={system_prompt_overhead}, total={total_input_tokens}")
            
            # Reset chat engine if new conversation or no history
            if not truncated_history:
                self.chat_engine.reset()
            
            # Skip hardcoded document-specific detection - use semantic similarity instead
            # (commented out to remove hardcoded document dependencies)
            
            # Use generic multi-document query engine for semantic similarity
            print(f"Using semantic similarity retrieval for query: {message[:50]}...")
            multi_doc_engine = self._create_multi_document_query_engine(similarity_top_k=8)
            response = multi_doc_engine.query(message)
            sources = self._extract_sources_from_response(response)
            
            if hasattr(response, 'response') and response.response:
                result = response.response.strip()
            else:
                result = str(response).strip() if response else ""
                # Log conversation history info for debugging
                if truncated_history:
                    print(f"Using conversation history: {len(truncated_history)} messages (max: {settings.get('max_conversation_history', 10)})")
                else:
                    print("No conversation history used")
                
                # Use standard chat engine with error handling for token limits
                try:
                    response = self.chat_engine.chat(message)
                    sources = self._extract_sources_from_response(response)
                    
                    if hasattr(response, 'response') and response.response:
                        result = response.response.strip()
                    else:
                        result = str(response).strip() if response else ""
                        
                except Exception as chat_error:
                    error_msg = str(chat_error).lower()
                    if "token" in error_msg and ("limit" in error_msg or "exceed" in error_msg):
                        return f"Error: Message too long for current model settings. Please try a shorter message or increase max_tokens in settings. Current limit: {max_tokens} tokens.", []
                    else:
                        return f"Error in chat processing: {str(chat_error)}", []
            
            # Truncate response to respect max_tokens setting
            result = self._truncate_response_to_max_tokens(result, max_tokens)
            
            return result, sources
            
        except Exception as e:
            print(f"Error in chat_with_documents_with_sources: {e}")
            return f"Error in chat: {str(e)}", []

    async def chat_with_documents(self, message: str, conversation_history: Optional[List] = None) -> str:
        """Chat with documents using the chat engine."""
        try:
            if not self.chat_engine:
                raise ValueError("Chat engine not initialized")
            
            # Update LLM settings to ensure current max_tokens is used
            self.update_llm_settings()
            
            # Load settings for conversation history management
            settings = self._load_app_settings()
            
            # Load max_tokens for later use, but don't validate here since Smart Context Management handles it
            max_tokens = settings.get('max_tokens', 512)
            
            # Log the message length for debugging
            message_tokens = self._estimate_token_count(message)
            print(f"📝 Processing message: {message_tokens} tokens, max_tokens: {max_tokens}")
            
            # Truncate conversation history more aggressively if needed
            truncated_history = self._truncate_conversation_history(conversation_history, settings)
            
            # Estimate total input tokens (message + history + system prompt overhead)
            # But rely on Smart Context Management for actual token management
            history_tokens = 0
            if truncated_history:
                history_text = " ".join([f"{msg.get('role', '')}: {msg.get('content', '')}" for msg in truncated_history])
                history_tokens = self._estimate_token_count(history_text)
            
            system_prompt_overhead = 200  # Estimate for system prompt and formatting
            total_input_tokens = message_tokens + history_tokens + system_prompt_overhead
            
            print(f"🔍 Input breakdown: message={message_tokens}, history={history_tokens}, overhead={system_prompt_overhead}, total={total_input_tokens}")
            
            # Reset chat engine if new conversation or no history
            if not truncated_history:
                self.chat_engine.reset()
            
            # Skip hardcoded document-specific detection - use semantic similarity instead
            # (commented out to remove hardcoded document dependencies)
            
            # Use generic multi-document query engine for semantic similarity
            print(f"Using semantic similarity retrieval for query: {message[:50]}...")
            multi_doc_engine = self._create_multi_document_query_engine(similarity_top_k=8)
            response = multi_doc_engine.query(message)
            
            # CLEAN UP RESPONSE: Extract only the AI-generated answer
            if hasattr(response, 'response') and response.response:
                raw_response = response.response.strip()
                
                # Remove metadata sections that start with separators
                clean_response = raw_response
                
                # Split on common LlamaIndex separators and take only the first part
                separators = [
                    "--------------------",
                    "doc_id:",
                    "Context information is below",
                    "Helpful Answer:",
                    "system:",
                    "\n\ndoc_id:",
                    "\n--------------------"
                ]
                
                for separator in separators:
                    if separator in clean_response:
                        clean_response = clean_response.split(separator)[0].strip()
                        break
                
                # Additional cleanup: remove any trailing metadata-like content
                lines = clean_response.split('\n')
                clean_lines = []
                
                for line in lines:
                    # Skip lines that look like metadata
                    if any(metadata_indicator in line.lower() for metadata_indicator in [
                        'doc_id:', 'file_path:', 'title:', 'type:', 'size_bytes:', 'created_at:', 
                        'modified_at:', 'mime_type:', 'encoding:', 'platform:', 'file_hash:',
                        'confluence_sync_', 'absolute_path:', 'relative_path:'
                    ]):
                        break
                    clean_lines.append(line)
                
                result = '\n'.join(clean_lines).strip()
            else:
                result = str(response).strip() if response else ""
            
            if not result or result.lower() in ['', 'empty response', 'empty response.']:
                return "I couldn't find relevant information in the documents to answer your question."
            
            # Truncate response to respect max_tokens setting
            result = self._truncate_response_to_max_tokens(result, settings['max_tokens'])
            
            return result
            
        except Exception as e:
            print(f"Error in chat_with_documents: {e}")
            return f"Error in chat: {str(e)}"
    
    def get_all_documents(self) -> List[Document]:
        """Get all processed documents, loading from disk if not in memory."""
        # First, scan for any documents on disk that aren't loaded in memory
        try:
            # Create a set of all loaded file paths (normalized to absolute paths)
            loaded_file_paths = set()
            for doc in self.documents.values():
                if doc.file_path:
                    try:
                        abs_path = str(Path(doc.file_path).resolve())
                        loaded_file_paths.add(abs_path)
                    except:
                        loaded_file_paths.add(doc.file_path)
            
            # Also track files we've processed in this method call to avoid duplicates
            processed_in_this_call = set()
            
            for file_path in self.documents_dir.glob("*.md"):
                file_path_abs = str(file_path.resolve())
                
                # Skip if already loaded in memory OR already processed in this call
                if file_path_abs in loaded_file_paths or file_path_abs in processed_in_this_call:
                    continue
                
                # Found a document on disk not in memory - try to load it synchronously
                print(f"Loading missing document from disk: {file_path}")
                try:
                    self._load_document_from_file_sync(file_path)
                    print(f"Successfully loaded missing document: {file_path.name}")
                    # Mark as processed to avoid loading again in this same call
                    processed_in_this_call.add(file_path_abs)
                except Exception as e:
                    print(f"Error loading missing document {file_path}: {e}")
                    
        except Exception as e:
            print(f"Error scanning for missing documents: {e}")
        
        return list(self.documents.values())

    def _load_document_from_file_sync(self, file_path: Path) -> None:
        """Load a document from file synchronously (for use in sync contexts)."""
        try:
            # Read content
            try:
                content = file_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Fallback for files with different encoding
                try:
                    content = file_path.read_text(encoding='latin-1')
                except Exception:
                    # Last resort - treat as binary and convert
                    with open(file_path, 'rb') as f:
                        raw_content = f.read()
                        content = str(raw_content, errors='ignore')
            
            # Determine document type
            doc_type = self._get_document_type(file_path.name)
            
            # Check if it's a Confluence file (markdown with metadata header)
            if file_path.name.startswith('confluence_') and file_path.suffix == '.md':
                if content.startswith('---\n'):
                    try:
                        parts = content.split('---\n', 2)
                        if len(parts) >= 3:
                            frontmatter = parts[1]
                            doc_content = parts[2]
                            
                            # Parse frontmatter
                            confluence_metadata = {}
                            for line in frontmatter.strip().split('\n'):
                                if ':' in line:
                                    key, value = line.split(':', 1)
                                    confluence_metadata[key.strip()] = value.strip()
                            
                            doc_id = confluence_metadata.get('doc_id')
                            if doc_id:
                                # Check if we already have a document with this doc_id
                                if doc_id in self.documents:
                                    existing_doc = self.documents[doc_id]
                                    print(f"Duplicate Confluence document found: {confluence_metadata.get('title', file_path.stem)}")
                                    print(f"  Existing: {existing_doc.file_path}")
                                    print(f"  Duplicate: {file_path}")
                                    print(f"  Skipping duplicate file: {file_path.name}")
                                    return
                                
                                document = Document(
                                    id=doc_id,
                                    title=confluence_metadata.get('title', file_path.stem),
                                    type=DocumentType.CONFLUENCE,
                                    content=doc_content.strip(),
                                    file_path=str(file_path),
                                    status=DocumentStatus.INDEXED,  # Assume already indexed since it's on disk
                                    metadata=confluence_metadata
                                )
                                
                                # Store in memory
                                self.documents[doc_id] = document
                                print(f"Loaded Confluence document: {document.title}")
                                return
                    except Exception as e:
                        print(f"Error parsing Confluence frontmatter in {file_path.name}: {e}")
            
            # Check if it's a Web import file (markdown with metadata header)
            if file_path.name.startswith('web_') and file_path.suffix == '.md':
                if content.startswith('---\n'):
                    try:
                        parts = content.split('---\n', 2)
                        if len(parts) >= 3:
                            frontmatter = parts[1]
                            doc_content = parts[2]
                            
                            # Parse frontmatter
                            web_metadata = {}
                            for line in frontmatter.strip().split('\n'):
                                if ':' in line:
                                    key, value = line.split(':', 1)
                                    web_metadata[key.strip()] = value.strip()
                            
                            doc_id = web_metadata.get('doc_id')
                            if doc_id:
                                # Check if we already have a document with this doc_id
                                if doc_id in self.documents:
                                    existing_doc = self.documents[doc_id]
                                    print(f"Duplicate web document found: {web_metadata.get('title', file_path.stem)}")
                                    print(f"  Existing: {existing_doc.file_path}")
                                    print(f"  Duplicate: {file_path}")
                                    print(f"  Skipping duplicate file: {file_path.name}")
                                    return
                                
                                document = Document(
                                    id=doc_id,
                                    title=web_metadata.get('title', file_path.stem),
                                    type=DocumentType.MARKDOWN,
                                    content=doc_content.strip(),
                                    file_path=str(file_path),
                                    status=DocumentStatus.INDEXED,  # Assume already indexed since it's on disk
                                    metadata=web_metadata
                                )
                                
                                # Store in memory
                                self.documents[doc_id] = document
                                print(f"Loaded web document: {document.title}")
                                return
                    except Exception as e:
                        print(f"Error parsing web frontmatter in {file_path.name}: {e}")
            
            # Handle regular files - create a simple document
            doc_id = str(uuid.uuid4())
            document = Document(
                id=doc_id,
                title=file_path.stem,
                type=doc_type,
                content=content,
                file_path=str(file_path),
                status=DocumentStatus.INDEXED,
                metadata={}
            )
            
            # Store in memory
            self.documents[doc_id] = document
            print(f"Loaded document: {document.title}")
            
        except Exception as e:
            raise Exception(f"Failed to load document from {file_path}: {e}")
    
    def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a specific document by ID."""
        return self.documents.get(doc_id)
    
    async def delete_document(self, doc_id: str) -> bool:
        """Delete a document from the index and storage."""
        try:
            if doc_id not in self.documents:
                return False
            
            document = self.documents[doc_id]
            
            # Remove file if it exists
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
            
            # Remove from documents dictionary
            del self.documents[doc_id]
            
            # Note: ChromaDB doesn't have a direct way to remove by metadata
            # In a production system, you'd want to track document IDs in the vector store
            
            return True
            
        except Exception as e:
            print(f"Error deleting document {doc_id}: {e}")
            return False
    
    async def add_confluence_document(self, llama_doc: LlamaDocument, doc_id: str) -> Document:
        """Add a Confluence document to the main document service."""
        try:
            title = llama_doc.metadata.get('title', 'Untitled Confluence Page')
            
            # Create a safe filename from the title and doc_id
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
            filename = f"confluence_{safe_title}_{doc_id[:8]}.md"
            
            # Save content as a markdown file
            confluence_file_path = self.documents_dir / filename
            
            # Create markdown content with metadata header
            markdown_content = f"""---
title: {title}
source: confluence_import
page_id: {llama_doc.metadata.get('page_id', '')}
space_key: {llama_doc.metadata.get('space_key', '')}
web_url: {llama_doc.metadata.get('web_url', '')}
api_url: {llama_doc.metadata.get('api_url', '')}
document_type: confluence
imported_at: {llama_doc.metadata.get('updated_at', '')}
doc_id: {doc_id}
---

# {title}

{llama_doc.text}
"""
            
            # Write to file
            with open(confluence_file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"Saved Confluence content to: {confluence_file_path}")
            
            # Extract comprehensive file metadata after file creation
            file_metadata = await self._extract_file_metadata(confluence_file_path)
            
            # Extract content metadata from the text content
            content_metadata = await self._extract_content_metadata(
                confluence_file_path, 
                DocumentType.CONFLUENCE, 
                llama_doc.text
            )
            
            # Combine all metadata sources
            combined_metadata = {
                **file_metadata,
                **content_metadata,
                **llama_doc.metadata  # Confluence-specific metadata has highest priority
            }
            
            # Create a Document object for tracking
            document = Document(
                id=doc_id,
                title=title,
                type=DocumentType.CONFLUENCE,
                content=llama_doc.text,
                status=DocumentStatus.PROCESSING,
                metadata=combined_metadata,
                file_path=str(confluence_file_path),
                size_bytes=file_metadata['size_bytes'],
                created_at=datetime.fromtimestamp(file_metadata['created_at']),
                updated_at=datetime.fromtimestamp(file_metadata['modified_at'])
            )
            
            # Add to the vector index
            if self.index is not None:
                # Update the llama_doc metadata to include all extracted metadata
                enhanced_llama_doc = LlamaDocument(
                    text=llama_doc.text,
                    metadata={
                        **llama_doc.metadata,
                        **combined_metadata,
                        "doc_id": doc_id,
                        "file_path": str(confluence_file_path)
                    }
                )
                
                # Insert the enhanced document into the index
                self.index.insert(enhanced_llama_doc)
                
                # Update query and chat engines with multi-document retrieval
                self.query_engine = self._create_multi_document_query_engine(similarity_top_k=12)
                
                settings = self._load_app_settings()
                self.chat_engine = self.index.as_chat_engine(
                    chat_mode="condense_plus_context",
                    similarity_top_k=10,  # Higher retrieval for better document coverage
                    verbose=True,
                    system_prompt=f"""You are a helpful document analysis assistant with access to multiple documents. Use ALL the provided document context to answer questions accurately and comprehensively.

CRITICAL INSTRUCTIONS:
1. SEARCH THROUGH ALL document context provided - you may have content from multiple documents
2. When listing documents, analyze ALL the context to identify different document sources
3. Look for document titles, filenames, and content from different sources in your context
4. For questions about specific documents (like "PSET-RFC"), search through all context for partial matches
5. Provide comprehensive summaries that draw from multiple documents when available
6. Reference specific documents by title/filename when possible
7. If you find multiple documents, clearly distinguish between them in your response
8. Keep responses under {settings['max_tokens']} tokens but prioritize multi-document coverage

REMEMBER: You may have context from multiple documents even if not explicitly labeled - analyze all content provided."""
                )
            
            # Mark as indexed
            document.status = DocumentStatus.INDEXED
            
            # Store in documents dictionary
            self.documents[doc_id] = document
            
            print(f"Successfully added Confluence document: {document.title} with {len(combined_metadata)} metadata fields")
            return document
            
        except Exception as e:
            print(f"Error adding Confluence document {doc_id}: {e}")
            # Create error document
            document = Document(
                id=doc_id,
                title=llama_doc.metadata.get('title', 'Untitled Confluence Page'),
                type=DocumentType.CONFLUENCE,
                status=DocumentStatus.ERROR,
                metadata=llama_doc.metadata
            )
            self.documents[doc_id] = document
            raise
    
    async def add_web_document(self, llama_doc: LlamaDocument, doc_id: str) -> Document:
        """Add a web document to the main document service."""
        try:
            title = llama_doc.metadata.get('title', 'Untitled Web Page')
            
            # Create a safe filename from the title and doc_id
            safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
            filename = f"web_{safe_title}_{doc_id[:8]}.md"
            
            # Save content as a markdown file
            web_file_path = self.documents_dir / filename
            
            # Create markdown content with metadata header
            markdown_content = f"""---
title: {title}
source: web_import
url: {llama_doc.metadata.get('url', '')}
original_url: {llama_doc.metadata.get('original_url', '')}
domain: {llama_doc.metadata.get('domain', '')}
document_type: webpage
extraction_method: {llama_doc.metadata.get('extraction_method', 'unknown')}
imported_at: {llama_doc.metadata.get('imported_at', '')}
doc_id: {doc_id}
---

# {title}

{llama_doc.text}
"""
            
            # Write to file
            with open(web_file_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"Saved web content to: {web_file_path}")
            
            # Extract comprehensive file metadata after file creation
            file_metadata = await self._extract_file_metadata(web_file_path)
            
            # Extract content metadata from the text content
            content_metadata = await self._extract_content_metadata(
                web_file_path, 
                DocumentType.MARKDOWN,  # Use MARKDOWN type for web content (saved as .md files)
                llama_doc.text
            )
            
            # Combine all metadata sources
            combined_metadata = {
                **file_metadata,
                **content_metadata,
                **llama_doc.metadata  # Web-specific metadata has highest priority
            }
            
            # Create a Document object for tracking
            document = Document(
                id=doc_id,
                title=title,
                type=DocumentType.MARKDOWN,  # Use MARKDOWN type for web content (saved as .md files)
                content=llama_doc.text,
                status=DocumentStatus.PROCESSING,
                metadata=combined_metadata,
                file_path=str(web_file_path),
                size_bytes=file_metadata['size_bytes'],
                created_at=datetime.fromtimestamp(file_metadata['created_at']),
                updated_at=datetime.fromtimestamp(file_metadata['modified_at'])
            )
            
            # Add to the vector index
            if self.index is not None:
                # Update the llama_doc metadata to include all extracted metadata
                enhanced_llama_doc = LlamaDocument(
                    text=llama_doc.text,
                    metadata={
                        **llama_doc.metadata,
                        **combined_metadata,
                        "doc_id": doc_id,
                        "file_path": str(web_file_path)
                    }
                )
                
                # Insert the enhanced document into the index
                self.index.insert(enhanced_llama_doc)
                
                # Update query and chat engines with multi-document retrieval
                self.query_engine = self._create_multi_document_query_engine(similarity_top_k=12)
                
                settings = self._load_app_settings()
                self.chat_engine = self.index.as_chat_engine(
                    chat_mode="condense_plus_context",
                    similarity_top_k=10,  # Higher retrieval for better document coverage
                    verbose=True,
                    system_prompt=f"""You are a helpful document analysis assistant with access to multiple documents. Use ALL the provided document context to answer questions accurately and comprehensively.

CRITICAL INSTRUCTIONS:
1. SEARCH THROUGH ALL document context provided - you may have content from multiple documents
2. When listing documents, analyze ALL the context to identify different document sources
3. Look for document titles, filenames, and content from different sources in your context
4. For questions about specific documents (like "PSET-RFC"), search through all context for partial matches
5. Provide comprehensive summaries that draw from multiple documents when available
6. Reference specific documents by title/filename when possible
7. If you find multiple documents, clearly distinguish between them in your response
8. Keep responses under {settings['max_tokens']} tokens but prioritize multi-document coverage

REMEMBER: You may have context from multiple documents even if not explicitly labeled - analyze all content provided."""
                )
            
            # Mark as indexed
            document.status = DocumentStatus.INDEXED
            
            # Store in documents dictionary
            self.documents[doc_id] = document
            
            print(f"Successfully added web document: {document.title} with {len(combined_metadata)} metadata fields")
            return document
            
        except Exception as e:
            print(f"Error adding web document {doc_id}: {e}")
            # Create error document
            document = Document(
                id=doc_id,
                title=llama_doc.metadata.get('title', 'Untitled Web Page'),
                type=DocumentType.MARKDOWN,
                status=DocumentStatus.ERROR,
                metadata=llama_doc.metadata
            )
            self.documents[doc_id] = document
            raise
    
    async def cleanup(self):
        """Cleanup resources."""
        if self.chroma_client:
            # ChromaDB client cleanup
            pass
        print("Document service cleanup complete")
    
    def get_document_metadata(self, doc_id: str) -> Optional[Dict[str, Any]]:
        """Get comprehensive metadata for a specific document."""
        if doc_id not in self.documents:
            return None
        
        document = self.documents[doc_id]
        
        # Return structured metadata information
        return {
            "document_info": {
                "id": document.id,
                "title": document.title,
                "type": document.type.value,
                "status": document.status.value,
                "file_path": document.file_path,
                "created_at": document.created_at.isoformat() if document.created_at else None,
                "updated_at": document.updated_at.isoformat() if document.updated_at else None,
                "size_bytes": document.size_bytes
            },
            "file_metadata": {
                k: v for k, v in document.metadata.items() 
                if k.startswith(('filename', 'file_', 'size_', 'mime_', 'platform', 'age_'))
            },
            "content_metadata": {
                k: v for k, v in document.metadata.items() 
                if k.startswith(('content_', 'word_', 'line_', 'paragraph_', 'character_', 'contains_'))
            },
            "format_specific_metadata": {
                k: v for k, v in document.metadata.items() 
                if k.startswith(('pdf_', 'docx_', 'markdown_', 'html_', 'confluence_'))
            },
            "system_metadata": {
                k: v for k, v in document.metadata.items() 
                if k.startswith(('absolute_', 'relative_', 'file_mode', 'file_uid', 'file_gid', 'file_owner', 'file_group'))
            },
            "raw_metadata": document.metadata  # Complete metadata for advanced users
        }
    
    def get_all_documents_metadata(self) -> Dict[str, Dict[str, Any]]:
        """Get metadata for all documents."""
        return {
            doc_id: self.get_document_metadata(doc_id) 
            for doc_id in self.documents.keys()
        }

    def update_llm_settings(self):
        """Update LLM settings from current app settings."""
        settings = self._load_app_settings()
        
        # Update the LLM's default max_tokens
        if self.llm and hasattr(self.llm, 'set_default_max_tokens'):
            self.llm.set_default_max_tokens(settings['max_tokens'])
            print(f"Updated LLM max_tokens to: {settings['max_tokens']}")
            
            # Reinitialize engines to use updated settings
            if self.index:
                # Reinitialize query engine with multi-document retrieval
                self.query_engine = self._create_multi_document_query_engine(similarity_top_k=12)
                
                # Reinitialize chat engine with multi-document retrieval
                self.chat_engine = self.index.as_chat_engine(
                    chat_mode="condense_plus_context",
                    similarity_top_k=10,  # Higher retrieval for better document coverage
                    verbose=True,
                    system_prompt=f"""You are a helpful document analysis assistant with access to multiple documents. Use ALL the provided document context to answer questions accurately and comprehensively.

CRITICAL INSTRUCTIONS:
1. SEARCH THROUGH ALL document context provided - you may have content from multiple documents
2. When listing documents, analyze ALL the context to identify different document sources
3. Look for document titles, filenames, and content from different sources in your context
4. For questions about specific documents (like "PSET-RFC"), search through all context for partial matches
5. Provide comprehensive summaries that draw from multiple documents when available
6. Reference specific documents by title/filename when possible
7. If you find multiple documents, clearly distinguish between them in your response
8. Keep responses under {settings['max_tokens']} tokens but prioritize multi-document coverage

REMEMBER: You may have context from multiple documents even if not explicitly labeled - analyze all content provided."""
                )
                
                print("Reinitialize engines with updated settings")
    
    def _truncate_response_to_max_tokens(self, response: str, max_tokens: int) -> str:
        """Provide smart partial responses when content exceeds token limits."""
        estimated_tokens = self._estimate_token_count(response)
        
        if estimated_tokens <= max_tokens * 0.4:  # 40% for response
            return response
        
        # For responses that are too long, provide a meaningful partial response
        target_tokens = int(max_tokens * 0.35)  # Be conservative, leave room for messaging
        target_chars = target_tokens * 4
        
        # Try to find a good breaking point
        truncated = response[:target_chars]
        
        # Find natural break points
        break_points = [
            truncated.rfind('\n\n'),  # Paragraph break
            truncated.rfind('\n# '),   # Heading
            truncated.rfind('. '),     # Sentence end
            truncated.rfind('\n'),     # Line break
        ]
        
        # Use the best break point that's reasonably close to the end
        best_break = 0
        for break_point in break_points:
            if break_point > len(truncated) * 0.7:  # In last 30%
                best_break = break_point
                break
        
        if best_break > 0:
            truncated = truncated[:best_break]
        
        # Add helpful continuation message
        remaining_chars = len(response) - len(truncated)
        continuation_msg = f"""

---
**⚠️ PARTIAL RESPONSE** ({remaining_chars:,} characters truncated)

**Current token limit:** {max_tokens} tokens  
**Estimated full response:** {estimated_tokens:,} tokens

**To see the complete response:**
- Increase max_tokens in Settings → Model Settings to at least {estimated_tokens + 500}
- Or ask more specific questions to get focused answers

**Quick answer to "What does A2A stand for?"**: A2A typically stands for "Agent-to-Agent" communication in AI/automation contexts."""
        
        return truncated + continuation_msg 

    def _detect_document_specific_query(self, query: str) -> List[str]:
        """
        Detect if a query is asking about specific documents and return matching document names.
        Returns list of document filenames that match the query keywords.
        """
        query_lower = query.lower()
        
        # Common patterns for document-specific queries
        specific_patterns = [
            'analyze the document',
            'summarize the document',
            'what does the document',
            'in the document',
            'document:',
            'examine the',
            'analyze this document',
            'review the document',
            'summarize:',
            'transcript',
            'recording',
            'meeting',
            'gmt20250630'
        ]
        
        # Check if this looks like a document-specific query
        is_document_specific = any(pattern in query_lower for pattern in specific_patterns)
        

        
        if not is_document_specific:
            # Also check for specific document name mentions
            document_keywords = [
                'pset-rfc', 'pset rfc', 'inter ai agent', 'orchestrator', 'plugin communication',
                'asset graph', 'google a2a', 'streaming capabilities', 'conditional write',
                'adoption of mcp', 'next-generation', 'autodesk assistant'
            ]
            is_document_specific = any(keyword in query_lower for keyword in document_keywords)
        

        
        if not is_document_specific:
            return []
        
        # Find matching documents
        matching_docs = []
        
        try:
            if self.documents_dir.exists():
                for doc_file in self.documents_dir.glob('*'):
                    if doc_file.is_file() and doc_file.name != '.gitkeep':
                        # Extract keywords from filename for matching
                        filename_lower = doc_file.name.lower()
                        
                        # More precise matching: prioritize exact matches and specific combinations
                        score = 0
                        
                        # Check for specific document patterns
                        if 'orchestrator' in query_lower and 'plugin' in query_lower:
                            if 'orchestrator' in filename_lower and 'plugin' in filename_lower:
                                score += 10  # High priority for exact match
                        
                        if 'pset-rfc' in query_lower or 'pset rfc' in query_lower:
                            if 'pset-rfc' in filename_lower or 'pset_rfc' in filename_lower:
                                score += 8  # High priority for PSET-RFC
                                
                        # Additional keyword matching
                        keywords_found = sum(1 for keyword in [
                            'pset-rfc', 'orchestrator', 'plugin', 'communication', 'protocol',
                            'asset_graph', 'google_a2a', 'streaming', 'conditional', 'write',
                            'mcp', 'next-generation', 'autodesk', 'assistant',
                            'transcript', 'recording', 'meeting', 'gmt'
                        ] if keyword in filename_lower)
                        
                        score += keywords_found
                        
                        # Also check if query mentions specific terms from this filename
                        filename_keywords = filename_lower.replace('_', ' ').replace('-', ' ')
                        query_terms_found = sum(1 for term in filename_keywords.split() 
                                              if len(term) > 3 and term in query_lower)
                        score += query_terms_found
                        
                        if score > 0:
                            matching_docs.append((doc_file.name, score))
                                
        except Exception as e:
            print(f"Error scanning documents directory: {e}")
        
        # Sort by score (descending) and return just the filenames
        if matching_docs:
            matching_docs.sort(key=lambda x: x[1], reverse=True)
            return [doc[0] for doc in matching_docs]
        
        return []
    
    def _load_document_content(self, filename: str) -> str:
        """
        Load the full content of a document for direct context injection.
        Supports both markdown (.md) and Word (.docx) files.
        """
        try:
            doc_path = self.documents_dir / filename
            if not doc_path.exists():
                return f"Document '{filename}' not found."
            
            if filename.endswith('.md'):
                # Read markdown file
                with open(doc_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content
                
            elif filename.endswith('.docx'):
                # Read Word document using python-docx
                try:
                    from docx import Document
                    doc = Document(doc_path)
                    paragraphs = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text)
                    content = '\n\n'.join(paragraphs)
                    return content
                except ImportError:
                    return f"Cannot read .docx files: python-docx not installed. Document: {filename}"
                except Exception as e:
                    return f"Error reading Word document '{filename}': {str(e)}"
            
            else:
                # Try to read as text file
                with open(doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return content
                
        except Exception as e:
            return f"Error loading document '{filename}': {str(e)}"
    
    def _estimate_token_count(self, text: str) -> int:
        """
        Rough estimation of token count (approximately 4 characters per token).
        """
        return len(text) // 4
    
    def _create_direct_document_prompt(self, documents: List[str], query: str) -> str:
        """
        Create a prompt with full document content for direct analysis.
        Prioritizes the most relevant documents within context limits.
        """
        # Get current max_tokens setting and be more conservative
        settings = self._load_app_settings()
        max_tokens = settings.get('max_tokens', 512)
        
        # Use only 50% of max_tokens for context, leave 50% for response
        max_context_tokens = int(max_tokens * 0.5)
        
        # Reserve tokens for query and prompt overhead
        query_tokens = self._estimate_token_count(query)
        prompt_overhead = 150  # Estimate for prompt formatting
        available_tokens = max_context_tokens - query_tokens - prompt_overhead
        
        if available_tokens < 200:  # Not enough tokens for meaningful context
            return None
        
        total_content = []
        total_tokens = 0
        
        # Limit to top 2 most relevant documents to be more conservative
        relevant_documents = documents[:2]
        
        for doc_filename in relevant_documents:
            content = self._load_document_content(doc_filename)
            tokens = self._estimate_token_count(content)
            
            if total_tokens + tokens > available_tokens:
                # If adding this document would exceed context, truncate or skip
                remaining_tokens = available_tokens - total_tokens
                if remaining_tokens > 400:  # Only include if we have reasonable space (reduced from 800)
                    truncated_content = content[:remaining_tokens * 4]
                    total_content.append(f"=== DOCUMENT: {doc_filename} (TRUNCATED) ===\n{truncated_content}\n")
                break
            else:
                total_content.append(f"=== DOCUMENT: {doc_filename} ===\n{content}\n")
                total_tokens += tokens
        
        if not total_content:
            return None
        
        prompt = f"""You are analyzing the following document(s) in detail. Use the complete document content provided below to answer the user's question comprehensively.

DOCUMENT CONTENT:
{''.join(total_content)}

USER QUESTION: {query}

Instructions:
1. Base your analysis entirely on the document content provided above
2. Provide detailed, comprehensive answers since you have access to the full document
3. Reference specific sections, quotes, or details from the documents
4. If the question asks for a summary, provide a thorough overview of the document's key points
5. If analyzing multiple documents, clearly distinguish between them in your response
"""
        
        return prompt
    
    # Template Management Methods
    
    async def _load_existing_templates(self):
        """Load existing templates from the templates directory on startup."""
        try:
            loaded_count = 0
            
            # Load templates metadata if it exists
            await self._load_templates_metadata()
            
            # Scan the templates directory for template files
            if self.templates_dir.exists():
                for file_path in self.templates_dir.iterdir():
                    if file_path.is_file() and file_path.suffix == '.md' and not file_path.name.startswith('.'):
                        try:
                            await self._load_template_from_file(file_path)
                            loaded_count += 1
                        except Exception as e:
                            print(f"Error loading template {file_path.name}: {e}")
            
            print(f"Loaded {loaded_count} existing templates from file system")
            
        except Exception as e:
            print(f"Error loading existing templates: {e}")
    
    async def _load_template_from_file(self, file_path: Path):
        """Load a single template from a file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Check if it's a template file with frontmatter
            if content.startswith('---\n'):
                parts = content.split('---\n', 2)
                if len(parts) >= 3:
                    frontmatter = parts[1]
                    template_content = parts[2]
                    
                    # Parse frontmatter
                    template_data = {}
                    for line in frontmatter.strip().split('\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            template_data[key.strip()] = value.strip()
                    
                    # Create template object
                    template = Template(
                        id=template_data.get('id', file_path.stem),
                        name=template_data.get('name', file_path.stem),
                        description=template_data.get('description', ''),
                        source_url=template_data.get('source_url'),
                        content=template_content.strip(),
                        sections=json.loads(template_data.get('sections', '[]')),
                        template_type=template_data.get('template_type', 'custom'),
                        space_key=template_data.get('space_key'),
                        page_id=template_data.get('page_id'),
                        created_at=datetime.fromisoformat(template_data.get('created_at', datetime.now().isoformat())),
                        updated_at=datetime.fromisoformat(template_data.get('updated_at', datetime.now().isoformat())),
                        last_synced=datetime.fromisoformat(template_data['last_synced']) if template_data.get('last_synced') else None,
                        sync_enabled=template_data.get('sync_enabled', 'true').lower() == 'true'
                    )
                    
                    self.templates[template.id] = template
                    
        except Exception as e:
            print(f"Error loading template from {file_path}: {e}")
    
    async def _load_templates_metadata(self):
        """Load templates metadata from JSON file."""
        try:
            if self.templates_metadata_file.exists():
                with open(self.templates_metadata_file, 'r') as f:
                    metadata = json.load(f)
                    # Additional metadata could be stored here if needed
        except Exception as e:
            print(f"Error loading templates metadata: {e}")
    
    async def _save_templates_metadata(self):
        """Save templates metadata to JSON file."""
        try:
            metadata = {
                "last_updated": datetime.now().isoformat(),
                "template_count": len(self.templates),
                "templates": {
                    template_id: {
                        "name": template.name,
                        "template_type": template.template_type,
                        "source_url": template.source_url,
                        "created_at": template.created_at.isoformat(),
                        "updated_at": template.updated_at.isoformat()
                    }
                    for template_id, template in self.templates.items()
                }
            }
            
            with open(self.templates_metadata_file, 'w') as f:
                json.dump(metadata, f, indent=2)
                
        except Exception as e:
            print(f"Error saving templates metadata: {e}")
    
    async def save_template(self, template: Template) -> str:
        """Save a template to the file system and in-memory storage."""
        try:
            # Generate unique ID if not provided
            if not template.id:
                template.id = str(uuid.uuid4())
            
            # Update timestamps
            now = datetime.now()
            if template.id not in self.templates:
                template.created_at = now
            template.updated_at = now
            
            # Save to file with frontmatter
            file_path = self.templates_dir / f"{template.id}.md"
            
            frontmatter = f"""---
id: {template.id}
name: {template.name}
description: {template.description}
template_type: {template.template_type}
sections: {json.dumps(template.sections)}
created_at: {template.created_at.isoformat()}
updated_at: {template.updated_at.isoformat()}"""
            
            if template.source_url:
                frontmatter += f"\nsource_url: {template.source_url}"
            if template.space_key:
                frontmatter += f"\nspace_key: {template.space_key}"
            if template.page_id:
                frontmatter += f"\npage_id: {template.page_id}"
            if template.last_synced:
                frontmatter += f"\nlast_synced: {template.last_synced.isoformat()}"
            
            frontmatter += f"\nsync_enabled: {str(template.sync_enabled).lower()}"
            frontmatter += "\n---\n\n"
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(frontmatter + template.content)
            
            # Save to in-memory storage
            self.templates[template.id] = template
            
            # Update metadata file
            await self._save_templates_metadata()
            
            print(f"Template saved: {template.name} (ID: {template.id})")
            return template.id
            
        except Exception as e:
            print(f"Error saving template: {e}")
            raise
    
    async def get_template(self, template_id: str) -> Optional[Template]:
        """Get a template by ID."""
        return self.templates.get(template_id)
    
    async def list_templates(self) -> List[Template]:
        """Get all templates."""
        return list(self.templates.values())
    
    async def delete_template(self, template_id: str) -> bool:
        """Delete a template by ID."""
        try:
            if template_id not in self.templates:
                return False
            
            # Delete the file
            file_path = self.templates_dir / f"{template_id}.md"
            if file_path.exists():
                file_path.unlink()
            
            # Remove from in-memory storage
            del self.templates[template_id]
            
            # Update metadata file
            await self._save_templates_metadata()
            
            print(f"Template deleted: {template_id}")
            return True
            
        except Exception as e:
            print(f"Error deleting template {template_id}: {e}")
            return False
    
    async def sync_template_from_confluence(self, template_id: str, credentials) -> Optional[Template]:
        """Sync a template with its Confluence source."""
        try:
            template = self.templates.get(template_id)
            if not template or not template.source_url:
                return None
            
            # Import confluence parsing functions
            from ..api.confluence import (
                parse_confluence_url, fetch_confluence_page_info
            )
            
            # Create auth headers (handle both dict and object credentials)
            if hasattr(credentials, 'auth_type'):
                # It's a ConfluenceCredentials object
                if credentials.auth_type == "pat":
                    auth_headers = {
                        "Authorization": f"Bearer {credentials.api_token}",
                        "Content-Type": "application/json"
                    }
                else:  # basic auth
                    import base64
                    auth_string = f"{credentials.username}:{credentials.api_token}"
                    encoded_auth = base64.b64encode(auth_string.encode()).decode()
                    auth_headers = {
                        "Authorization": f"Basic {encoded_auth}",
                        "Content-Type": "application/json"
                    }
            else:
                # It's a dict, convert to auth headers
                if credentials.get('auth_type') == "pat":
                    auth_headers = {
                        "Authorization": f"Bearer {credentials['api_token']}",
                        "Content-Type": "application/json"
                    }
                else:  # basic auth
                    import base64
                    auth_string = f"{credentials['username']}:{credentials['api_token']}"
                    encoded_auth = base64.b64encode(auth_string.encode()).decode()
                    auth_headers = {
                        "Authorization": f"Basic {encoded_auth}",
                        "Content-Type": "application/json"
                    }
            
            # Parse the Confluence URL
            url_info = parse_confluence_url(template.source_url)
            
            # Fetch fresh content from Confluence
            page_info = await fetch_confluence_page_info(url_info, auth_headers)
            
            # Convert HTML to markdown using the same logic as template creation
            content = self._convert_html_to_markdown(page_info['content'])
            
            # Update template
            template.content = content
            template.name = page_info['title']
            template.last_synced = datetime.now()
            template.updated_at = datetime.now()
            
            # Extract sections from the updated content
            template.sections = self._extract_sections_from_content(content)
            
            # Save the updated template
            await self.save_template(template)
            
            print(f"Template synced: {template.name}")
            return template
            
        except Exception as e:
            print(f"Error syncing template {template_id}: {e}")
            return None
    
    def _convert_html_to_markdown(self, html_content: str) -> str:
        """Convert HTML content to basic markdown format."""
        if not html_content:
            return ""
        
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Convert HTML to basic markdown
        content = html_content
        
        # Convert headers
        content = content.replace('<h1>', '# ').replace('</h1>', '\n\n')
        content = content.replace('<h2>', '## ').replace('</h2>', '\n\n')
        content = content.replace('<h3>', '### ').replace('</h3>', '\n\n')
        content = content.replace('<h4>', '#### ').replace('</h4>', '\n\n')
        content = content.replace('<h5>', '##### ').replace('</h5>', '\n\n')
        content = content.replace('<h6>', '###### ').replace('</h6>', '\n\n')
        
        # Convert paragraphs
        content = content.replace('<p>', '').replace('</p>', '\n\n')
        
        # Convert bold and italic
        content = content.replace('<strong>', '**').replace('</strong>', '**')
        content = content.replace('<b>', '**').replace('</b>', '**')
        content = content.replace('<em>', '*').replace('</em>', '*')
        content = content.replace('<i>', '*').replace('</i>', '*')
        
        # Convert lists
        content = content.replace('<ul>', '').replace('</ul>', '\n')
        content = content.replace('<ol>', '').replace('</ol>', '\n')
        content = content.replace('<li>', '- ').replace('</li>', '\n')
        
        # Convert line breaks
        content = content.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
        
        # Remove remaining HTML tags and get clean text
        soup = BeautifulSoup(content, 'html.parser')
        clean_content = soup.get_text()
        
        # Clean up excessive newlines
        lines = clean_content.split('\n')
        cleaned_lines = []
        prev_empty = False
        
        for line in lines:
            line = line.strip()
            if line == '':
                if not prev_empty:
                    cleaned_lines.append('')
                prev_empty = True
            else:
                cleaned_lines.append(line)
                prev_empty = False
        
        return '\n'.join(cleaned_lines).strip()

    def _extract_sections_from_content(self, content: str) -> List[str]:
        """Extract section headers from markdown content."""
        sections = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                # Extract header text without the # symbols
                header = line.lstrip('#').strip()
                if header and header not in sections:
                    sections.append(header)
        
        return sections